#Signing the code
#File Created By Héctor Rodríguez Fusté
#SysAdmin / OS & Software Engineer
#Code for the Sox Form File & Access Form File
#Version: 1 - Created on 19/01/2017
ProgramVersion = "4.0.0"
DateVersion = "19/12/2023"

#importing the libraries
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
from openpyxl import load_workbook
from openpyxl import Workbook
import shutil, os, time, logging, win32com.client
from easygui import *

wdFormatPDF = 17
Ordinal = ["Unknown","st","nd","rd","th","th","th","th","th","th","th","st","nd","rd","th","th","th","th","th","th","th","st","nd","rd","th","th","th","th","th","th","th","st","nd","rd","th","th","th","th","th","th","th","st","nd","rd","th","th","th","th","th","th","th","st","nd","rd","th","th","th","th","th","th","th"]


def covx_to_pdf(infile):
    """Convert a Word .docx to PDF"""
    pwd = os.getcwd().replace("\\","\\\\")

    infile = pwd + "\\" + infile
    outfile = infile.replace(".docx",".pdf")
    # word = comtypes.client.CreateObject('Word.Application')
    word = win32com.client.Dispatch('Word.Application')
    doc = word.Documents.Open(infile)
    doc.SaveAs(outfile, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()
def charCounter(string = None, char = " "):
    var = len(string) - len(string.replace(char,""))
    return var
def htmlContent(modification = False, SoxFormFile = False):
   with open(".\\Files\\DefaultSoxEmailTemplate.html", "r", encoding="utf-8") as htmlfile:
       DefaultSoxEmail=htmlfile.read()

   with open(".\\Files\\DefaultModifiedSoxEmailTemplate.html", "r", encoding="utf-8") as htmlfile:
       DefaultModifiedSoxEmail=htmlfile.read()

   with open(".\\Files\\DefaultAccessFormEmailTemplate.html", "r", encoding="utf-8") as htmlfile:
       DefaultAccessFormEmail=htmlfile.read()

   if modification == False:
       return DefaultSoxEmail
   elif modification == True:
       return DefaultModifiedSoxEmail
   if SoxFormFile == True:
       return DefaultAccessFormEmail
def sendEmail(NameFile, addr_to = None, addr_cc = None, addr_from = None, FormDayDate = None, ModifyFile = False, SubjectMsg = None, PdfFile = None, NameAttach = None, SendSoxForm = False, UserName = None, StartDate = None, Dept = None):
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.base import MIMEBase
    from email import encoders
    import getpass
    Username = "UserName"

    if addr_to == None and addr_cc == None and addr_from == None:
        addr_to   = 'Company_ISM@Company.com'
        addr_cc   = 'Company_HRD@Company.com;'
        rcpt = addr_cc.split(";") + [addr_to]
        addr_from = 'Company.HRD@Company.com'

    else:
        addr_to   = addr_to
        addr_cc   = addr_cc
        rcpt = addr_cc.split(";") + [addr_to]
        addr_from = addr_from

    ##el mail sale desde el correo
    smtp_server = 'smtpmail.Company.com'

    msg = MIMEMultipart('alternative')
    msg['To'] = addr_to
    msg['Cc'] = addr_cc
    msg['From'] = addr_from
    if ModifyFile == False:
        msg['Subject'] = "Sox Form " + str(FormDayDate) + str(Ordinal[int(FormDayDate)]) + " " + MonthName + " "
        msg.attach(MIMEText(htmlContent(),'html'))
    elif ModifyFile == True:
        msg['Subject'] = "Sox Form " + str(FormDayDate) + str(Ordinal[int(FormDayDate)]) + " " + MonthName + " - " + "Modified: " + str(DayNumber) + "." + MonthNumber
        msg.attach(MIMEText(htmlContent(modification = True),'html'))

    if SendSoxForm == True:
        msg['Subject'] = "User " + Username + " - " + "Start Date: " + StartDate + " - " + Dept
        msg.attach(MIMEText(htmlContent(SoxFormFile = True),'html'))
    
    #cuerpo del mensaje en HTML y si fuera solo text puede colocar en el 2da parametro 'plain'

    #adjuntamos fichero de texto pero puede ser cualquer tipo de archivo
    ##cargamos el archivo a adjuntar
    if ModifyFile == False:
        NameAttach = "Company Sox" + " - " + str(FormDayDate) + str(Ordinal[int(FormDayDate)]) + " " + MonthName + " - " + str(YearNumber) + ".docx"
    elif ModifyFile == True:
        NameAttach = "Company Sox" + " - " + str(FormDayDate) + str(Ordinal[int(FormDayDate)]) + " " + MonthName + " - " + str(YearNumber) + " - Modified" + ".docx"

    if SendSoxForm == True:
        NameAttach = PdfFile

    fp = open(NameFile,'rb')
    adjunto = MIMEBase('multipart', 'encrypted')
    #lo insertamos en una variable
    adjunto.set_payload(fp.read())
    fp.close()
    #lo encriptamos en base64 para enviarlo
    encoders.encode_base64(adjunto)
    #agregamos una cabecera y le damos un nombre al archivo que adjuntamos puede ser el mismo u otro
    adjunto.add_header('Content-Disposition', 'attachment', filename=NameAttach)
    #adjuntamos al mensaje
    msg.attach(adjunto)

    # inicializamos el stmp para hacer el envio
    server = smtplib.SMTP(smtp_server,25)
    server.starttls()
    #logeamos con los datos ya seteamos en la parte superior
    # server.login(smtp_user,smtp_pass)
    #el envio
    server.sendmail(addr_from, rcpt, msg.as_string())
    #apagamos conexion stmp
    server.quit()
    return
def terminatingUsers(FirstName, LastName, EmployeeNum, LeavingDate, TerminatedNameFile):
    docW = Document("Files\\TerminatedForm.docx")
    TerminatedUserName = FirstName + " " + LastName
     
    docW._body.clear_content()
    ###TABLE CREATION###
    EHeaderTable = docW.add_table(rows=2,cols=4)
    # p = docW.add_paragraph(" ")
    docW.add_paragraph().add_run().add_break()
    EBodyTable = docW.add_table(rows=18,cols=5)
    # p = docW.add_paragraph(" ")
    docW.add_paragraph().add_run().add_break()
    ESignsTable = docW.add_table(rows=2,cols=6)
    docW.add_paragraph().add_run().add_break()
    # EUserSignTable = docW.add_table(rows=1,cols=1, style = "Table Sign")
    # EUserSignTable = docW.add_table(rows=1,cols=1, style = "Table Grid")
    EUserTable = docW.add_table(rows=1,cols=1, style="Table Sign")

    ###PUTTING DATA INTO THE HEADER TABLE###
    EHeaderTable.rows[0].cells[0].text = "LAST NAME:"
    EHeaderTable.rows[0].cells[1].text = LastName
    EHeaderTable.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(14)
    EHeaderTable.rows[0].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    EHeaderTable.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    EHeaderTable.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    EHeaderTable.rows[0].cells[0].paragraphs[0].runs[0].bold = True

    EHeaderTable.rows[0].cells[2].text = "FIRST NAME:"
    EHeaderTable.rows[0].cells[3].text = FirstName
    EHeaderTable.rows[0].cells[2].paragraphs[0].runs[0].font.size = Pt(14)
    EHeaderTable.rows[0].cells[3].paragraphs[0].runs[0].font.size = Pt(11)
    EHeaderTable.rows[0].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    EHeaderTable.rows[0].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    EHeaderTable.rows[0].cells[2].paragraphs[0].runs[0].bold = True

    EHeaderTable.rows[1].cells[0].text = "PAYROLL EMPLOYEE NUM:"
    EHeaderTable.rows[1].cells[1].text = EmployeeNum
    EHeaderTable.rows[1].cells[0].paragraphs[0].runs[0].font.size = Pt(14)
    EHeaderTable.rows[1].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    EHeaderTable.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    EHeaderTable.rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    EHeaderTable.rows[1].cells[0].paragraphs[0].runs[0].bold = True

    EHeaderTable.rows[1].cells[2].text = "END DATE:"
    EHeaderTable.rows[1].cells[3].text = LeavingDate
    EHeaderTable.rows[1].cells[2].paragraphs[0].runs[0].font.size = Pt(14)
    EHeaderTable.rows[1].cells[3].paragraphs[0].runs[0].font.size = Pt(11)
    EHeaderTable.rows[1].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    EHeaderTable.rows[1].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    EHeaderTable.rows[1].cells[2].paragraphs[0].runs[0].bold = True

    ###PUTTING DATA INTO THE BODY TABLE###
    ###PUTTING DATA INTO THE BODY TABLE###
    ###PUTTING DATA INTO THE BODY TABLE###
    EBodyTable.rows[0].cells[0].text = "DEPT01"
    EBodyTable.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(12)
    EBodyTable.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    EBodyTable.rows[0].cells[0].paragraphs[0].runs[0].underline = True
    EBodyTable.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    EBodyTable.rows[1].cells[0].text = "EXAMPLE01"
    EBodyTable.rows[1].cells[1].text = "RETURNED"
    EBodyTable.rows[1].cells[2].text = "YES"
    EBodyTable.rows[1].cells[3].text = "NO"
    EBodyTable.rows[1].cells[4].text = "N/A"
    # EBodyTable.rows[1].cells[2].add_paragraph(text=("YES").replace(" ",""), style="Check List")
    # EBodyTable.rows[1].cells[3].add_paragraph(text="NO", style="Check List")
    EBodyTable.rows[1].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    EBodyTable.rows[1].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    EBodyTable.rows[1].cells[2].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[1].cells[3].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[1].cells[4].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    EBodyTable.rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    EBodyTable.rows[1].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[1].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[1].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[1].cells[0].paragraphs[0].runs[0].italic = True

    EBodyTable.rows[2].cells[0].text = "HUMAN RESOURCES"
    EBodyTable.rows[2].cells[0].paragraphs[0].runs[0].font.size = Pt(12)
    EBodyTable.rows[2].cells[0].paragraphs[0].runs[0].bold = True
    EBodyTable.rows[2].cells[0].paragraphs[0].runs[0].underline = True
    EBodyTable.rows[2].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    a=EBodyTable.rows[2].cells[0]
    b=EBodyTable.rows[2].cells[1]

    a.merge(b)

    EBodyTable.rows[3].cells[0].text = "EXAMPLE02"
    EBodyTable.rows[3].cells[1].text = "RETURNED"
    EBodyTable.rows[3].cells[2].text = "YES"
    EBodyTable.rows[3].cells[3].text = "NO"
    EBodyTable.rows[3].cells[4].text = "N/A"
    # EBodyTable.rows[3].cells[2].add_run().List(style="Check List")
    # EBodyTable.rows[3].cells[2].add_paragraph(text="YES", style="Check List")
    # EBodyTable.rows[3].cells[3].add_paragraph(text="NO", style="Check List")
    EBodyTable.rows[3].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    EBodyTable.rows[3].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    EBodyTable.rows[3].cells[2].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[3].cells[3].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[3].cells[4].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[3].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    EBodyTable.rows[3].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    EBodyTable.rows[3].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[3].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[3].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[3].cells[0].paragraphs[0].runs[0].italic = True
    # 
    EBodyTable.rows[4].cells[0].text = "PAYROLL SYSTEM"
    EBodyTable.rows[4].cells[1].text = "TERMINATED"
    EBodyTable.rows[4].cells[2].text = "YES"
    EBodyTable.rows[4].cells[3].text = "NO"
    EBodyTable.rows[4].cells[4].text = "N/A"
    # EBodyTable.rows[4].cells[2].add_paragraph(text="YES", style="Check List")
    # EBodyTable.rows[4].cells[3].add_paragraph(text="NO", style="Check List")
    EBodyTable.rows[4].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    EBodyTable.rows[4].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    EBodyTable.rows[4].cells[2].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[4].cells[3].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[4].cells[4].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[4].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    EBodyTable.rows[4].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    EBodyTable.rows[4].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[4].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[4].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[4].cells[0].paragraphs[0].runs[0].italic = True

    EBodyTable.rows[5].cells[0].text = "TIME ATTENDANCE SYSTEM"
    EBodyTable.rows[5].cells[1].text = "TERMINATED"
    EBodyTable.rows[5].cells[2].text = "YES"
    EBodyTable.rows[5].cells[3].text = "NO"
    EBodyTable.rows[5].cells[4].text = "N/A"
    # EBodyTable.rows[5].cells[2].add_paragraph(text="YES", style="Check List")
    # EBodyTable.rows[5].cells[3].add_paragraph(text="NO", style="Check List")
    EBodyTable.rows[5].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    EBodyTable.rows[5].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    EBodyTable.rows[5].cells[2].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[5].cells[3].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[5].cells[4].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[5].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    EBodyTable.rows[5].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    EBodyTable.rows[5].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[5].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[5].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[5].cells[0].paragraphs[0].runs[0].italic = True

    EBodyTable.rows[6].cells[0].text = "DEPT02"
    EBodyTable.rows[6].cells[0].paragraphs[0].runs[0].font.size = Pt(12)
    EBodyTable.rows[6].cells[0].paragraphs[0].runs[0].bold = True
    EBodyTable.rows[6].cells[0].paragraphs[0].runs[0].underline = True
    EBodyTable.rows[6].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    EBodyTable.rows[7].cells[0].text = "EXAMPLE03"
    EBodyTable.rows[7].cells[1].text = "RETURNED"
    EBodyTable.rows[7].cells[2].text = "YES"
    EBodyTable.rows[7].cells[3].text = "NO"
    EBodyTable.rows[7].cells[4].text = "N/A"
    # EBodyTable.rows[7].cells[2].add_paragraph(text="YES", style="Check List")
    # EBodyTable.rows[7].cells[3].add_paragraph(text="NO", style="Check List")
    EBodyTable.rows[7].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    EBodyTable.rows[7].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    EBodyTable.rows[7].cells[2].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[7].cells[3].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[4].cells[4].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[7].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    EBodyTable.rows[7].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    EBodyTable.rows[7].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[7].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[7].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[7].cells[0].paragraphs[0].runs[0].italic = True

    EBodyTable.rows[8].cells[0].text = "DEPT03"
    EBodyTable.rows[8].cells[0].paragraphs[0].runs[0].font.size = Pt(12)
    EBodyTable.rows[8].cells[0].paragraphs[0].runs[0].bold = True
    EBodyTable.rows[8].cells[0].paragraphs[0].runs[0].underline = True
    EBodyTable.rows[8].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    EBodyTable.rows[9].cells[0].text = "EXAMPLE04"
    EBodyTable.rows[9].cells[1].text = "RETURNED"
    EBodyTable.rows[9].cells[2].text = "YES"
    EBodyTable.rows[9].cells[3].text = "NO"
    EBodyTable.rows[9].cells[4].text = "N/A"
    # EBodyTable.rows[9].cells[2].add_paragraph(text="YES", style="Check List")
    # EBodyTable.rows[9].cells[3].add_paragraph(text="NO", style="Check List")
    EBodyTable.rows[9].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    EBodyTable.rows[9].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    EBodyTable.rows[9].cells[2].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[9].cells[3].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[9].cells[4].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[9].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    EBodyTable.rows[9].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    EBodyTable.rows[9].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[9].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[9].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[9].cells[0].paragraphs[0].runs[0].italic = True

    EBodyTable.rows[10].cells[0].text = "EXAMPLE05"
    EBodyTable.rows[10].cells[1].text = "TERMINATED"
    EBodyTable.rows[10].cells[2].text = "YES"
    EBodyTable.rows[10].cells[3].text = "NO"
    EBodyTable.rows[10].cells[4].text = "N/A"
    # EBodyTable.rows[10].cells[2].add_paragraph(text="YES", style="Check List")
    # EBodyTable.rows[10].cells[3].add_paragraph(text="NO", style="Check List")
    EBodyTable.rows[10].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    EBodyTable.rows[10].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    EBodyTable.rows[10].cells[2].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[10].cells[3].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[10].cells[4].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[10].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    EBodyTable.rows[10].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    EBodyTable.rows[10].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[10].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[10].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[10].cells[0].paragraphs[0].runs[0].italic = True

    EBodyTable.rows[11].cells[0].text = "INFORMATION TECHNOLOGY"
    EBodyTable.rows[11].cells[0].paragraphs[0].runs[0].font.size = Pt(12)
    EBodyTable.rows[11].cells[0].paragraphs[0].runs[0].bold = True
    EBodyTable.rows[11].cells[0].paragraphs[0].runs[0].underline = True
    EBodyTable.rows[11].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    a = EBodyTable.rows[11].cells[0]
    b = EBodyTable.rows[11].cells[1]
    a.merge(b)

    EBodyTable.rows[12].cells[0].text = "EXAMPLE06"
    EBodyTable.rows[12].cells[1].text = "RETURNED"
    EBodyTable.rows[12].cells[2].text = "YES"
    EBodyTable.rows[12].cells[3].text = "NO"
    EBodyTable.rows[12].cells[4].text = "N/A"
    # EBodyTable.rows[12].cells[2].add_paragraph(text="YES", style="Check List")
    # EBodyTable.rows[12].cells[3].add_paragraph(text="NO", style="Check List")
    EBodyTable.rows[12].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    EBodyTable.rows[12].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    EBodyTable.rows[12].cells[2].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[12].cells[3].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[12].cells[4].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[12].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    EBodyTable.rows[12].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    EBodyTable.rows[12].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[12].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[12].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[12].cells[0].paragraphs[0].runs[0].italic = True

    EBodyTable.rows[13].cells[0].text = "EXAMPLE07"
    EBodyTable.rows[13].cells[1].text = "RETURNED"
    EBodyTable.rows[13].cells[2].text = "YES"
    EBodyTable.rows[13].cells[3].text = "NO"
    EBodyTable.rows[13].cells[4].text = "N/A"
    # EBodyTable.rows[13].cells[2].add_paragraph(text="YES", style="Check List")
    # EBodyTable.rows[13].cells[3].add_paragraph(text="NO", style="Check List")
    EBodyTable.rows[13].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    EBodyTable.rows[13].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    EBodyTable.rows[13].cells[2].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[13].cells[3].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[13].cells[4].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[13].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    EBodyTable.rows[13].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    EBodyTable.rows[13].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[13].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[13].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[13].cells[0].paragraphs[0].runs[0].italic = True

    EBodyTable.rows[14].cells[0].text = "EXAMPLE08"
    EBodyTable.rows[14].cells[1].text = "RETURNED"
    EBodyTable.rows[14].cells[2].text = "YES"
    EBodyTable.rows[14].cells[3].text = "NO"
    EBodyTable.rows[14].cells[4].text = "N/A"
    # EBodyTable.rows[14].cells[2].add_paragraph(text="YES", style="Check List")
    # EBodyTable.rows[14].cells[3].add_paragraph(text="NO", style="Check List")
    EBodyTable.rows[14].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    EBodyTable.rows[14].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    EBodyTable.rows[14].cells[2].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[14].cells[3].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[14].cells[4].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[14].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    EBodyTable.rows[14].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    EBodyTable.rows[14].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[14].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[14].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[14].cells[0].paragraphs[0].runs[0].italic = True

    EBodyTable.rows[15].cells[0].text = "ACTIVE DIRECTORY"
    EBodyTable.rows[15].cells[1].text = "TERMINATED"
    EBodyTable.rows[15].cells[2].text = "YES"
    EBodyTable.rows[15].cells[3].text = "NO"
    EBodyTable.rows[15].cells[4].text = "N/A"
    # EBodyTable.rows[15].cells[2].add_paragraph(text="YES", style="Check List")
    # EBodyTable.rows[15].cells[3].add_paragraph(text="NO", style="Check List")
    EBodyTable.rows[15].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    EBodyTable.rows[15].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    EBodyTable.rows[15].cells[2].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[15].cells[3].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[15].cells[4].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[15].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    EBodyTable.rows[15].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    EBodyTable.rows[15].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[15].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[15].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[15].cells[0].paragraphs[0].runs[0].italic = True

    EBodyTable.rows[16].cells[0].text = "DEPT04"
    EBodyTable.rows[16].cells[0].paragraphs[0].runs[0].font.size = Pt(12)
    EBodyTable.rows[16].cells[0].paragraphs[0].runs[0].bold = True
    EBodyTable.rows[16].cells[0].paragraphs[0].runs[0].underline = True
    EBodyTable.rows[16].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    a = EBodyTable.rows[16].cells[0]
    b = EBodyTable.rows[16].cells[1]
    a.merge(b)

    EBodyTable.rows[17].cells[0].text = "EXAMPLE09"
    EBodyTable.rows[17].cells[1].text = "RETURNED"
    EBodyTable.rows[17].cells[2].text = "YES"
    EBodyTable.rows[17].cells[3].text = "NO"
    EBodyTable.rows[17].cells[4].text = "N/A"
    # EBodyTable.rows[17].cells[2].add_paragraph(text="YES", style="Check List")
    # EBodyTable.rows[17].cells[3].add_paragraph(text="NO", style="Check List")
    EBodyTable.rows[17].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    EBodyTable.rows[17].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    EBodyTable.rows[17].cells[2].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[17].cells[3].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[17].cells[4].paragraphs[0].runs[0].font.size = Pt(10)
    EBodyTable.rows[17].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    EBodyTable.rows[17].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    EBodyTable.rows[17].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[17].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[17].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EBodyTable.rows[17].cells[0].paragraphs[0].runs[0].italic = True


    ###PUTTING DATA INTO THE FOOTER TABLE###
    ###PUTTING DATA INTO THE FOOTER TABLE###
    ###PUTTING DATA INTO THE FOOTER TABLE###
    ###PUTTING DATA INTO THE FOOTER TABLE###

    ESignsTable.rows[0].cells[0].text = "SIGNATURE & DATE FOR DEPT01 SUPERVISOR"
    ESignsTable.rows[0].cells[1].text = "SIGNATURE & DATE FOR HUMAN RESOURCES SUPERVISOR"
    ESignsTable.rows[0].cells[2].text = "SIGNATURE & DATE FOR DEPT02 SUPERVISOR"
    ESignsTable.rows[0].cells[3].text = "SIGNATURE & DATE FOR DEPT03 SUPERVISOR"
    ESignsTable.rows[0].cells[4].text = "SIGNATURE & DATE FOR INFORMATION TECHNOLOGY SUPERVISOR"
    ESignsTable.rows[0].cells[5].text = "SIGNATURE & DATE FOR DEPT04 SUPERVISOR"
    ESignsTable.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(12)
    ESignsTable.rows[0].cells[1].paragraphs[0].runs[0].font.size = Pt(12)
    ESignsTable.rows[0].cells[2].paragraphs[0].runs[0].font.size = Pt(12)
    ESignsTable.rows[0].cells[3].paragraphs[0].runs[0].font.size = Pt(12)
    ESignsTable.rows[0].cells[4].paragraphs[0].runs[0].font.size = Pt(12)
    ESignsTable.rows[0].cells[5].paragraphs[0].runs[0].font.size = Pt(12)
    ESignsTable.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    ESignsTable.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    ESignsTable.rows[0].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    ESignsTable.rows[0].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    ESignsTable.rows[0].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    ESignsTable.rows[0].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    ESignsTable.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    ESignsTable.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    ESignsTable.rows[0].cells[2].paragraphs[0].runs[0].bold = True
    ESignsTable.rows[0].cells[3].paragraphs[0].runs[0].bold = True
    ESignsTable.rows[0].cells[4].paragraphs[0].runs[0].bold = True
    ESignsTable.rows[0].cells[5].paragraphs[0].runs[0].bold = True

    ESignsTable.rows[1].cells[0].text = " "
    ESignsTable.rows[1].cells[1].text = " "
    ESignsTable.rows[1].cells[2].text = " "
    ESignsTable.rows[1].cells[3].text = " "
    ESignsTable.rows[1].cells[4].text = " "
    ESignsTable.rows[1].cells[5].text = " "
    ESignsTable.rows[1].cells[0].paragraphs[0].runs[0].add_break()

    EUserTable.rows[0].cells[0].text = str(TerminatedUserName)
    EUserTable.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(12)
    EUserTable.rows[0].cells[0].paragraphs[0].runs[0].bold = True

    # EUserTable.rows[0].cells[0].paragraphs[0].runs[0].add_break()
    # EUserTable.rows[0].cells[0].paragraphs[0].runs[0].add_break()
    # EUserTable.rows[0].cells[0].paragraphs[0].runs[0].add_break()

    docW.save(TerminatedNameFile)
    covx_to_pdf(TerminatedNameFile)
def replaceSpecialChars(String):
    z = 0
    for j in specialChar:
        if charCounter(string = String, char = j) > 0:
            if z == 0:
                String = String.replace(str(j),correctChar[0])
            elif z == 1:
                String = String.replace(str(j),correctChar[0])
            elif z == 2:
                String = String.replace(str(j),correctChar[0])
            elif z == 3:
                String = String.replace(str(j),correctChar[1])
            elif z == 4:
                String = String.replace(str(j),correctChar[1])
            elif z == 5:
                String = String.replace(str(j),correctChar[1])
            elif z == 6:
                String = String.replace(str(j),correctChar[2])
            elif z == 7:
                String = String.replace(str(j),correctChar[2])
            elif z == 8:
                String = String.replace(str(j),correctChar[2])
            elif z == 9:
                String = String.replace(str(j),correctChar[3])
            elif z == 10:
                String = String.replace(str(j),correctChar[3])
            elif z == 11:
                String = String.replace(str(j),correctChar[3])
            elif z == 12:
                String = String.replace(str(j),correctChar[4])
            elif z == 13:
                String = String.replace(str(j),correctChar[4])
            elif z == 14:
                String = String.replace(str(j),correctChar[4])
            elif z == 15:
                String = String.replace(str(j),correctChar[5])
            elif z == 16:
                String = String.replace(str(j),correctChar[6])
            elif z == 17:
                String = String.replace(str(j),"")
            elif z == 18:
                String = String.replace(str(j),correctChar[7])
        z+=1

    return String
def copySoxData(SoxFile, AccessFile):
    fromCopy = load_workbook(SoxFile) #File where we have the data saved.
    toCopy = load_workbook(AccessFile) #File where we have the data saved.
    Sox = fromCopy["Sox"] #The name of the data sheet.
    Access = toCopy["Sox"] #The name of the data sheet.

    for x in range(2,(Sox.max_row)+1):
        Access.cell(row=x,column=1, value = Sox.cell(row=x,column=1).value)
        Access.cell(row=x,column=2, value = Sox.cell(row=x,column=2).value)
        Access.cell(row=x,column=3, value = Sox.cell(row=x,column=3).value)
        Access.cell(row=x,column=4, value = Sox.cell(row=x,column=4).value)
        Access.cell(row=x,column=5, value = Sox.cell(row=x,column=5).value)
        Access.cell(row=x,column=6, value = Sox.cell(row=x,column=6).value)
        Access.cell(row=x,column=7, value = Sox.cell(row=x,column=7).value)
        Access.cell(row=x,column=8, value = Sox.cell(row=x,column=8).value)
        Access.cell(row=x,column=9, value = Sox.cell(row=x,column=9).value)
        Access.cell(row=x,column=10,value = Sox.cell(row=x,column=10).value)
    else:
        toCopy.save(AccessFile)
        docE = load_workbook(AccessFile)
        Sox = docE["Sox"]
        Data = docE["Users"]

os.makedirs("logs", exist_ok=True)
logfile = "logs\\CreateSoxForm-" + time.strftime("%H%M-%A_%d_%B_%Y") + ".log"

logging.basicConfig(level=logging.INFO, filemode='w', filename=logfile)

try:
    months = ["Unknown","January","February","March","April","May","June","July","August","September","October","November","December"]
    specialChar = ["á","à","ä","é","è","ë","í","ì","ï","ó","ò","ö","ú","ù","ü","ñ","ç","'","-"]
    correctChar = ["a","e","i","o","u","n","c"," "]
    #Variables for name folder.
    OutputFiles = "Sox Files"
    FormFiles = OutputFiles + "\\" + "Access Form"
    TableFiles = OutputFiles + "\\" + "Sox Form"

    # create folders, no error if it already exists
    MonthName= time.strftime("%B")
    MonthNumber= time.strftime("%m")
    YearNumber = time.strftime("%Y")
    DayNumber = time.strftime("%d")
    os.makedirs(OutputFiles, exist_ok=True)
    os.makedirs(FormFiles, exist_ok=True)
    os.makedirs(TableFiles, exist_ok=True)
    os.makedirs(FormFiles + "\\TMP", exist_ok=True)
    os.makedirs(TableFiles + "\\TMP", exist_ok=True)
    os.makedirs(TableFiles + "\\" + YearNumber, exist_ok=True)
    os.makedirs(FormFiles + "\\" + YearNumber, exist_ok=True)
    os.makedirs(TableFiles + "\\" + YearNumber + "\\" + MonthName, exist_ok=True)
    os.makedirs(FormFiles + "\\" + YearNumber + "\\" + MonthName, exist_ok=True)
    os.makedirs(FormFiles + "\\" + YearNumber + "\\" + MonthName + "\\Terminated", exist_ok=True)

    #Landscape or Portrait Mode
    SoxFormFile = "Files\\UserListForm.docx"
    # AccessFormFile = "Files\\NetworkUserFormLand.docx"
    AccessFormFile = "Files\\NetworkUserForm.docx"
    TerminatedFormFile = "Files\\TerminatedForm.docx"

    docW = Document(SoxFormFile) #Template for the Word file

    #FIRST WINDOW
    ProgramTitle = "Sox File Creator - By: Hector Rodriguez Fuste - Assistant IT Manager"
    MainMenuMsg = "Program Created By Hector Rodriguez\nAssistant ISM - Company Diagonal Mar Barcelona\nProgram Version: " + ProgramVersion + "\nDate Version: " + DateVersion + "\n\nWhat would you like to do? Click a button to choose an option"
    MainMenuChoices= ["[1] Make Sox Form", "[2] Modify a Sox Form", "[3] Make Access Form", "[4] Exit"]
    MainMenuChoice = "5"

    #OPTIONS FOR FIRST WINDOW
    AddRowSoxFileMsg = "How many Users will you add into the new Sox File?"
    DataSoxFileFields = ["First Name", "Last Name", "Personal Email", "Site Code", "Time Attendance Number", "Employee Number", "Position", "Date"]
    ModifyDataSoxFileFields = ["First Name", "Last Name", "Personal Email", "Site Code", "Time Attendance Number", "Employee Number", "Position", "Department", "Date"]
    DefaultSoxData = ["First Name", "Last Name", "Personal.Email@Company.com", "INNCODE", "XXXX", "YYYY", "Position", "DD/MM/YYYY"]
    DepartmentList = ["DEPT01", "DEPT02", "DEPT03", "DEPT04", "DEPT05"]
    #GUI SETUP FOR 2ND OPTION (MODIFY A ROW)
    ModifyUserSoxFile = ["User First Name", "User Last Name"]

    #MODIFY SOX FORM OPTION SETUP
    ModifySoxFileFields = ["Day of Sox File Creation", "Month of Sox File Creation", "Year of Sox File Creation"]
    DefaultValuesDataSoxFile = ["DD", "M", "YYYY"]
    ModifySoxMenuChoices = ["[1] Add a new row", "[2] Modify a row", "[3] Delete a row", "[4] Remake File", "[5] Back"]
    DuplicateFileMsg = "I detect 2 files. Which one would you like to modify?\n"

    #GUI SETUP FOR ACCESS FORM OPTION
    AccessOptionsMsg = "Choose an option to perform\n"
    AccessOptionsChoices = ["[1] Make Access Form", "[2] Resend Emails", "[3] Recreate Terminated Form", "[4] Copy Data from Sox File", "[5] Exit"]
    AccessFormFileFields = ["Network Access", "Intranet Access", "Program01 Access", "Program01 Rights", "Program02 Access", "Program03 Access", "Program04 Access", "Company Email Access", "Program05 Access", "Program06 Access", "PayRoll Access", "Time Attendance Access", "Program07 Access", "Program08 Access", "Program09 Access", "Key01 Access"]

    Exit = "No"

    while Exit == "No":

        MainMenuChoice = ((buttonbox(msg=MainMenuMsg, title=ProgramTitle, choices=MainMenuChoices)).split("] ")[0]).replace("[","")
        if MainMenuChoice == "1":
            docE = load_workbook("Files\\ExcelTemplate.xlsx") #File where we have the data saved.
            Data = docE["Users"] #The name of the data sheet.
            Sox = docE["Sox"] #The name of the data sheet.

            NumUsers = int(enterbox(msg=AddRowSoxFileMsg, title=ProgramTitle, default="1"))

            for x in range(1,int(NumUsers)+1):
                DefaultSoxData[7] = time.strftime("%d/%m/%Y")
                AddingDataToSoxFile = multenterbox(msg="Enter The Employee Data", title=ProgramTitle, fields=DataSoxFileFields, values=DefaultSoxData)
                Dept = choicebox(title=ProgramTitle, msg="Choose the Employee Department for " + "\n\nUsername: " + AddingDataToSoxFile[0] + " " + AddingDataToSoxFile[1] + "\nPosition: " + AddingDataToSoxFile[6], choices=DepartmentList)
                StartLeave = ((buttonbox(msg="Is the user \"Starting or Leaving\"", title=ProgramTitle, choices=["[1] Starting", "[2] Leaving", "[3] Cancel"])).split("] ")[0]).replace("[","")

                FirstName = AddingDataToSoxFile[0]
                FirstName = FirstName.strip()
                FirstName = FirstName.capitalize()
                LastName = AddingDataToSoxFile[1]
                LastName = LastName.strip()
                LastName = LastName.capitalize()

                if len(FirstName) == 0:
                    FirstName = "N/A"
                if charCounter(FirstName) == 1:
                    FirstName = FirstName.split(" ")[0].capitalize() + " " +  FirstName.split(" ")[1].capitalize()
                elif charCounter(FirstName) == 2:
                    FirstName = FirstName.split(" ")[0].capitalize() + " " +  FirstName.split(" ")[1].capitalize() + " " +  FirstName.split(" ")[2].capitalize()
                elif charCounter(FirstName) == 3:
                    FirstName = FirstName.split(" ")[0].capitalize() + " " +  FirstName.split(" ")[1].capitalize() + " " +  FirstName.split(" ")[2].capitalize() + " " +  FirstName.split(" ")[3].capitalize()

                if len(LastName) == 0:
                    LastName = "N/A"
                if charCounter(LastName) == 1:
                    LastName = LastName.split(" ")[0].capitalize() + " " +  LastName.split(" ")[1].capitalize()
                elif charCounter(LastName) == 2:
                    LastName = LastName.split(" ")[0].capitalize() + " " +  LastName.split(" ")[1].capitalize() + " " +  LastName.split(" ")[2].capitalize()
                elif charCounter(LastName) == 3:
                    LastName = LastName.split(" ")[0].capitalize() + " " +  LastName.split(" ")[1].capitalize() + " " +  LastName.split(" ")[2].capitalize() + " " +  LastName.split(" ")[3].capitalize()

                FirstName = replaceSpecialChars(FirstName)
                LastName = replaceSpecialChars(LastName)

                if AddingDataToSoxFile[2] == "Personal.Email@Company.com" or AddingDataToSoxFile[2] == "personal.email@Company.com" or len(AddingDataToSoxFile[2]) == 0:
                    PersonalEmail = "-"
                else:
                    PersonalEmail = AddingDataToSoxFile[2]

                OtherCompany = AddingDataToSoxFile[3]
                OtherCompany = OtherCompany.strip()
                TimeAttendanceNum = AddingDataToSoxFile[4]
                EmployeeNum = AddingDataToSoxFile[5]
                Position = AddingDataToSoxFile[6]
                Position = Position.strip()
                Position = Position.capitalize()

                if StartLeave == "1":
                    StartLeave = "Starting"
                elif StartLeave == "2":
                    StartLeave = "Leaving"
                else:
                    break

                SLDate = AddingDataToSoxFile[7]

                """ if len(PersonalEmail) == 0 or PersonalEmail == DefaultSoxData[2]:
                     PersonalEmail = "-"
                 else:
                     PersonalEmail = PersonalEmail.strip()
                """

                if len(OtherCompany) == 0:
                    OtherCompany = "-"
                if OtherCompany != DefaultSoxData[3]:
                    OtherCompany = OtherCompany.upper()
                else:
                    OtherCompany = "-"

                if len(TimeAttendanceNum) == 0 or TimeAttendanceNum == "XXXX":
                    TimeAttendanceNum = "-"

                if len(EmployeeNum) == 0 or EmployeeNum == "YYYY":
                    EmployeeNum = "-"

                if len(Position) == 0:
                    Position = "N/A"
                if charCounter(Position) == 1:
                    Position = Position.split(" ")[0].capitalize() + " " +  Position.split(" ")[1].capitalize()
                elif charCounter(Position) == 2:
                    Position = Position.split(" ")[0].capitalize() + " " +  Position.split(" ")[1].capitalize() + " " +  Position.split(" ")[2].capitalize()


                if len(SLDate) == 0 or SLDate =="DD/MM/YYYY":
                    SLDate = time.strftime("%d/%m/%Y")

                Sox.cell(row=x+1,column=1, value=FirstName)
                Sox.cell(row=x+1,column=2, value=LastName)
                Sox.cell(row=x+1,column=3, value=PersonalEmail)
                Sox.cell(row=x+1,column=4, value=OtherCompany)
                Sox.cell(row=x+1,column=5, value=TimeAttendanceNum)
                Sox.cell(row=x+1,column=6, value=EmployeeNum)
                Sox.cell(row=x+1,column=7, value=Position)
                Sox.cell(row=x+1,column=8, value=Dept)
                Sox.cell(row=x+1,column=9, value=StartLeave)
                Sox.cell(row=x+1,column=10, value=SLDate)
            else:
                SoxName = "Company Sox " + MonthName + " - " + str(YearNumber) + " - " + str(int(DayNumber))
                SoxNameModified = "Company Modified Sox " + MonthName + " - " + str(YearNumber) + " - " + str(int(DayNumber))

                TMPExcel = TableFiles + "\\TMP\\" + SoxName + ".xlsx"
                TMPExcelModified = TableFiles + "\\TMP\\" + SoxNameModified + ".xlsx"
                
            ReplaceOption = False
            print(os.path.exists(TMPExcel))
            print(os.path.isfile(TMPExcel))

            if os.path.exists(TMPExcel) == True and os.path.isfile(TMPExcel) == True:
                ReplaceOption = boolbox(msg="The file you want to create already exists. Do you want to create a Modified File ? If you click \"No\", the actual Sox File will be overwritten")
                if ReplaceOption == True:
                    print("Soy el reemplazo")
                    docE.save(TMPExcelModified)
                if ReplaceOption == False:
                    print("Soy el original")
                    docE.save(TMPExcel)
            else:
                docE.save(TMPExcel)


            
            if ReplaceOption == False:
                docE = load_workbook(TMPExcel) #File where we have the data saved.
            else:
                docE = load_workbook(TMPExcelModified) #File where we have the data saved.

            Data = docE["Users"] #The name of the data sheet.
            Sox = docE["Sox"] #The name of the data sheet.
            docW = Document(SoxFormFile)

            docW._body.clear_content()
            #User information Table
            HeaderTable = docW.add_table(rows=1, cols=1)
            Templatetable = docW.add_table(rows=1, cols=10, style="Table Grid")

            """
            FIRST TABLE/FIRST TABLE/FIRST TABLE
            """
            HeaderTable.rows[0].cells[0].text = "Sox Form Date: " + str(time.strftime("%d/%m/%Y"))
            HeaderTable.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(14)
            HeaderTable.rows[0].cells[0].paragraphs[0].runs[0].bold = True
            HeaderTable.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            Templatetable.rows[0].cells[0].text = "First Name"
            Templatetable.rows[0].cells[1].text = "Last Name"
            Templatetable.rows[0].cells[2].text = "Personal e-mail"
            Templatetable.rows[0].cells[3].text = "Site Code"
            Templatetable.rows[0].cells[4].text = "Time Attendance Employee Nr."
            Templatetable.rows[0].cells[5].text = "PayRoll Employee Nr."
            Templatetable.rows[0].cells[6].text = "Position"
            Templatetable.rows[0].cells[7].text = "Department"
            Templatetable.rows[0].cells[8].text = "Starting / Leaving (Alta / Baja)"
            Templatetable.rows[0].cells[9].text = "Date"
            Templatetable.rows[0].cells[0].paragraphs[0].runs[0].bold = True
            Templatetable.rows[0].cells[1].paragraphs[0].runs[0].bold = True
            Templatetable.rows[0].cells[2].paragraphs[0].runs[0].bold = True
            Templatetable.rows[0].cells[3].paragraphs[0].runs[0].bold = True
            Templatetable.rows[0].cells[4].paragraphs[0].runs[0].bold = True
            Templatetable.rows[0].cells[5].paragraphs[0].runs[0].bold = True
            Templatetable.rows[0].cells[6].paragraphs[0].runs[0].bold = True
            Templatetable.rows[0].cells[7].paragraphs[0].runs[0].bold = True
            Templatetable.rows[0].cells[8].paragraphs[0].runs[0].bold = True
            Templatetable.rows[0].cells[9].paragraphs[0].runs[0].bold = True
            Templatetable.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
            Templatetable.rows[0].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
            Templatetable.rows[0].cells[2].paragraphs[0].runs[0].font.size = Pt(11)
            Templatetable.rows[0].cells[3].paragraphs[0].runs[0].font.size = Pt(11)
            Templatetable.rows[0].cells[4].paragraphs[0].runs[0].font.size = Pt(11)
            Templatetable.rows[0].cells[5].paragraphs[0].runs[0].font.size = Pt(11)
            Templatetable.rows[0].cells[6].paragraphs[0].runs[0].font.size = Pt(11)
            Templatetable.rows[0].cells[7].paragraphs[0].runs[0].font.size = Pt(11)
            Templatetable.rows[0].cells[8].paragraphs[0].runs[0].font.size = Pt(11)
            Templatetable.rows[0].cells[9].paragraphs[0].runs[0].font.size = Pt(11)
            Templatetable.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            Templatetable.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            Templatetable.rows[0].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            Templatetable.rows[0].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            Templatetable.rows[0].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            Templatetable.rows[0].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            Templatetable.rows[0].cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            Templatetable.rows[0].cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            Templatetable.rows[0].cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            Templatetable.rows[0].cells[9].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            LineDeleted = 1

            for x in range(2,Sox.max_row+1):

                FirstName = Sox.cell(row=x,column=1).value
                LastName = Sox.cell(row=x,column=2).value
                PersonalEmail = Sox.cell(row=x,column=3).value
                OtherCompany = Sox.cell(row=x,column=4).value
                TimeAttendanceNum = Sox.cell(row=x,column=5).value
                EmployeeNum = Sox.cell(row=x,column=6).value
                Position = Sox.cell(row=x,column=7).value
                Dept = Sox.cell(row=x,column=8).value
                StartLeave = Sox.cell(row=x,column=9).value
                SLDate = Sox.cell(row=x,column=10).value

                NumLines = x - LineDeleted

                Templatetable.add_row()
                Templatetable.rows[NumLines].cells[0].text = FirstName
                Templatetable.rows[NumLines].cells[1].text = LastName
                Templatetable.rows[NumLines].cells[2].text = PersonalEmail
                Templatetable.rows[NumLines].cells[3].text = OtherCompany
                Templatetable.rows[NumLines].cells[4].text = TimeAttendanceNum
                Templatetable.rows[NumLines].cells[5].text = EmployeeNum
                Templatetable.rows[NumLines].cells[6].text = Position
                Templatetable.rows[NumLines].cells[7].text = Dept
                Templatetable.rows[NumLines].cells[8].text = StartLeave
                Templatetable.rows[NumLines].cells[9].text = SLDate
                Templatetable.rows[NumLines].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
                Templatetable.rows[NumLines].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
                Templatetable.rows[NumLines].cells[2].paragraphs[0].runs[0].font.size = Pt(11)
                Templatetable.rows[NumLines].cells[3].paragraphs[0].runs[0].font.size = Pt(11)
                Templatetable.rows[NumLines].cells[4].paragraphs[0].runs[0].font.size = Pt(11)
                Templatetable.rows[NumLines].cells[5].paragraphs[0].runs[0].font.size = Pt(11)
                Templatetable.rows[NumLines].cells[6].paragraphs[0].runs[0].font.size = Pt(11)
                Templatetable.rows[NumLines].cells[7].paragraphs[0].runs[0].font.size = Pt(11)
                Templatetable.rows[NumLines].cells[8].paragraphs[0].runs[0].font.size = Pt(11)
                if StartLeave == "Leaving":
                    Templatetable.rows[NumLines].cells[8].paragraphs[0].runs[0].font.color.rgb = RGBColor(200, 0, 0)
                else:
                    Templatetable.rows[NumLines].cells[8].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 200, 0)
                Templatetable.rows[NumLines].cells[9].paragraphs[0].runs[0].font.size = Pt(11)
                Templatetable.rows[NumLines].cells[8].paragraphs[0].runs[0].bold = True
                Templatetable.rows[NumLines].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                Templatetable.rows[NumLines].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                Templatetable.rows[NumLines].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                Templatetable.rows[NumLines].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                Templatetable.rows[NumLines].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                Templatetable.rows[NumLines].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                Templatetable.rows[NumLines].cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                Templatetable.rows[NumLines].cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                Templatetable.rows[NumLines].cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                Templatetable.rows[NumLines].cells[9].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                if StartLeave == "Leaving":

                    DateYForm = SLDate.split("/")[2]
                    DateMForm = SLDate.split("/")[1]
                    CorrectFormDate = SLDate.replace("/","-")
                    os.makedirs(FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)], exist_ok=True)
                    os.makedirs(FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated", exist_ok=True)

                    if charCounter(string = FirstName, char = " ") == 1:
                        if charCounter(string = LastName, char = " ") == 1:
                            NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                        elif charCounter(string = LastName, char = " ") == 2:
                            NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "-" + CorrectFormDate + ".docx"
                        elif charCounter(string = LastName, char = " ") == 3:
                            NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "_" + LastName.split(" ")[3] + "-" + CorrectFormDate + ".docx"
                        else:
                            NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName + "-" + CorrectFormDate + ".docx"
                    elif charCounter(string = FirstName, char = " ") == 2:
                        if charCounter(string = LastName, char = " ") == 1:
                            NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                        elif charCounter(string = LastName, char = " ") == 2:
                            NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "-" + CorrectFormDate + ".docx"
                        elif charCounter(string = LastName, char = " ") == 3:
                            NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "_" + LastName.split(" ")[3] + "-" + CorrectFormDate + ".docx"
                        else:
                            NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName + "-" + CorrectFormDate + ".docx"
                    elif charCounter(string = FirstName, char = " ") == 3:
                        if charCounter(string = LastName, char = " ") == 1:
                            NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                        elif charCounter(string = LastName, char = " ") == 2:
                            NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "-" + CorrectFormDate + ".docx"
                        elif charCounter(string = LastName, char = " ") == 3:
                            NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "_" + LastName.split(" ")[3] + "-" + CorrectFormDate + ".docx"
                        else:
                            NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName + "-" + CorrectFormDate + ".docx"

                    else:
                        if charCounter(string = LastName, char = " ") == 1:
                            NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                        else:
                            NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName + "_" + LastName + "-" + CorrectFormDate + ".docx"

                    terminatingUsers(FirstName, LastName, EmployeeNum, SLDate, NameTerminatedFile)
            else:
                SoxName = "Company Sox " + MonthName + " - " + str(YearNumber) + " - " + str(int(DayNumber))
                SoxNameModified = "Company Modified Sox " + MonthName + " - " + str(YearNumber) + " - " + str(int(DayNumber))

                NameFile = TableFiles + "\\" + YearNumber + "\\" + MonthName + "\\" + SoxName + ".docx"
                NameFileModified = TableFiles + "\\" + YearNumber + "\\" + MonthName + "\\" + SoxNameModified + ".docx"

            if os.path.exists(NameFile) == True and os.path.isfile(NameFile) == True:
                docW.save(NameFileModified)
            else:
                docW.save(NameFile)

            askForSend = boolbox(msg="Do you want to send this file by Email?")
            if askForSend == True:
                if ReplaceOption == False:
                    sendEmail(NameFile=NameFile, FormDayDate=str(int(DayNumber)))
                elif ReplaceOption == True:
                    sendEmail(NameFile=NameFileModified, FormDayDate=str(int(DayNumber)))
        elif MainMenuChoice == "2":
            FileFound = False
            FailedCountNum = 0
            

            while FileFound == False or FailedCountNum > 3: #Getting the file to modify
                DateSoxFile = multenterbox(msg="Enter Data for Sox Form File", title=ProgramTitle, fields=ModifySoxFileFields, values=DefaultValuesDataSoxFile)

                FileDay = DateSoxFile[0]
                FileMonth = DateSoxFile[1]
                FileYear = DateSoxFile[2]
                TMPExcel1 = "Company Sox " + str(months[int(FileMonth)]) + " - " + FileYear + " - " + FileDay
                TMPExcel2 = "Company Modified Sox " + str(months[int(FileMonth)]) + " - " + FileYear + " - " + FileDay
                TMPExcel = TableFiles + "\\TMP" + "\\" + TMPExcel1 + ".xlsx"
                TMPExcelModified = TableFiles + "\\TMP" + "\\" + TMPExcel2 + ".xlsx"
                DuplicateFileChoices = ["[1] " + TMPExcel1, "[2] " + TMPExcel2]

                if os.path.exists(TMPExcel) == True and os.path.exists(TMPExcelModified) == True:
                    optionModify = ((buttonbox(msg=DuplicateFileMsg, title=ProgramTitle, choices=DuplicateFileChoices)).split("] ")[0]).replace("[","")
                    if optionModify == "2":
                        docE = load_workbook(TMPExcelModified) #File where we have the data saved.
                        Data = docE["Users"] #The name of the data sheet.
                        Sox = docE["Sox"] #The name of the data sheet.
                        ModifySoxMenuMsg = "You're going to modify the sox file called: " + TMPExcel2 + "\n\n\nWhat would you like to do? Click a button to choose an option"
                        FileFound = True
                    else:
                        docE = load_workbook(TMPExcel) #File where we have the data saved.
                        Data = docE["Users"] #The name of the data sheet.
                        Sox = docE["Sox"] #The name of the data sheet.
                        ModifySoxMenuMsg = "You're going to modify the sox file called: " + TMPExcel1 + "\n\n\nWhat would you like to do? Click a button to choose an option"
                        FileFound = True
                elif os.path.exists(TMPExcel) == True and os.path.exists(TMPExcelModified) == False:
                    docE = load_workbook(TMPExcel) #File where we have the data saved.
                    Data = docE["Users"] #The name of the data sheet.
                    Sox = docE["Sox"] #The name of the data sheet.
                    ModifySoxMenuMsg = "You're going to modify the sox file called: " + TMPExcel1 + "\n\n\nWhat would you like to do? Click a button to choose an option"
                    ModifyFileMsg = "You're going to modify the sox file called: " + TMPExcel1 + "\n\n\nDo you want to modify this file, or make a \"modified\" file?"
                    ModifingFile = buttonbox(title=ProgramTitle, msg=ModifyFileMsg, choices=["Modify File","Make \"Modified\""])
                    if ModifingFile == "2":
                        optionModify = "2"
                    else:
                        optionModify = "1"
                    FileFound = True

                if FileFound == False:
                    DateSoxFile =  msgbox(msg="Error. File not found. Try it again!", title=ProgramTitle ,ok_button="Okay")
                    FailedCountNum+=1

            AnswerOption = ((buttonbox(msg=ModifySoxMenuMsg, title=ProgramTitle, choices=ModifySoxMenuChoices)).split("] ")[0]).replace("[","")
            if AnswerOption == "1":
                NumUsers = int(enterbox(msg=AddRowSoxFileMsg, title=ProgramTitle, default="1"))

                for x in range(2,((Sox.max_row + NumUsers)+1)):

                    if x > Sox.max_row:
                        DefaultSoxData[7] = SLDate = time.strftime("%d/%m/%Y")
                        AddingDataToSoxFile = multenterbox(msg="Enter The Employee Data", title=ProgramTitle, fields=DataSoxFileFields, values=DefaultSoxData)
                        Dept = choicebox(title=ProgramTitle, msg="Choose the Employee Department for " + "\n\nUsername: " + AddingDataToSoxFile[0] + " " + AddingDataToSoxFile[1] + "\nPosition: " + AddingDataToSoxFile[6], choices=DepartmentList)
                        StartLeave = ((buttonbox(msg="Is the user \"Starting or Leaving\"", title=ProgramTitle, choices=["[1] Starting", "[2] Leaving", "[3] Cancel"])).split("] ")[0]).replace("[","")

                        FirstName = AddingDataToSoxFile[0]
                        LastName = AddingDataToSoxFile[1]
                        if AddingDataToSoxFile[2] == "Personal.Email@Company.com" or AddingDataToSoxFile[2] == "personal.email@Company.com":
                            PersonalEmail = "-"
                        else:
                            PersonalEmail = AddingDataToSoxFile[2]
                        OtherCompany = AddingDataToSoxFile[3]
                        TimeAttendanceNum = AddingDataToSoxFile[4]
                        EmployeeNum = AddingDataToSoxFile[5]
                        Position = AddingDataToSoxFile[6]
                        SLDate = AddingDataToSoxFile[7]

                        FirstName = FirstName.capitalize()
                        LastName = LastName.capitalize()
                        PersonalEmail = PersonalEmail.capitalize()
                        OtherCompany = OtherCompany.upper()
                        Position = Position.capitalize()

                        if len(FirstName) == 0:
                            FirstName = "N/A"
                        elif charCounter(FirstName) == 1:
                            FirstName = FirstName.split(" ")[0].capitalize() + " " +  FirstName.split(" ")[1].capitalize()
                        elif charCounter(FirstName) == 2:
                            FirstName = FirstName.split(" ")[0].capitalize() + " " +  FirstName.split(" ")[1].capitalize() + " " +  FirstName.split(" ")[2].capitalize()

                        if len(LastName) == 0:
                            LastName = "N/A"
                        elif charCounter(LastName) == 1:
                            LastName = LastName.split(" ")[0].capitalize() + " " +  LastName.split(" ")[1].capitalize()
                        elif charCounter(LastName) == 2:
                            LastName = LastName.split(" ")[0].capitalize() + " " +  LastName.split(" ")[1].capitalize() + " " +  LastName.split(" ")[2].capitalize()
                        
                        FirstName = replaceSpecialChars(FirstName)
                        LastName = replaceSpecialChars(LastName)


                        if len(OtherCompany) == 0:
                            OtherCompany = "-"
                        elif OtherCompany != DefaultSoxData[3]:
                            OtherCompany = OtherCompany.upper()
                        else:
                            OtherCompany = "-"

                        if len(TimeAttendanceNum) == 0:
                            TimeAttendanceNum = "-"

                        if len(EmployeeNum) == 0:
                            EmployeeNum = "-"

                        if len(Position) == 0:
                            Position = "-"
                        elif charCounter(Position) == 1:
                            Position = Position.split(" ")[0].capitalize() + " " +  Position.split(" ")[1].capitalize()
                        elif charCounter(Position) == 2:
                            Position = Position.split(" ")[0].capitalize() + " " +  Position.split(" ")[1].capitalize() + " " +  Position.split(" ")[2].capitalize()

                        if StartLeave == "1":
                            StartLeave = "Starting"
                        elif StartLeave == "2":
                            StartLeave = "Leaving"
                        else:
                            msgbox(msg="Starting will be used as default Starting/Leaving data")
                            StartLeave = "Starting"

                        if len(SLDate) == 0 or SLDate =="DD/MM/YYYY":
                            SLDate = DefaultSoxData[7]

                        Sox.cell(row=x,column=1, value=FirstName)
                        Sox.cell(row=x,column=2, value=LastName)
                        Sox.cell(row=x,column=3, value=PersonalEmail)
                        Sox.cell(row=x,column=4, value=OtherCompany)
                        Sox.cell(row=x,column=5, value=TimeAttendanceNum)
                        Sox.cell(row=x,column=6, value=EmployeeNum)
                        Sox.cell(row=x,column=7, value=Position)
                        Sox.cell(row=x,column=8, value=Dept)
                        Sox.cell(row=x,column=9, value=StartLeave)
                        Sox.cell(row=x,column=10, value=SLDate)
                    else:
                        continue
                else:
                    if optionModify == "1":
                        docE.save(TMPExcel)
                        docE = load_workbook(TMPExcel) #File where we have the data saved.
                        Data = docE["Users"] #The name of the data sheet.
                        Sox = docE["Sox"] #The name of the data sheet.
                    elif optionModify == "2":
                        docE.save(TMPExcelModified) #File where we have the data saved.
                        docE = load_workbook(TMPExcelModified) #File where we have the data saved.
                        Data = docE["Users"] #The name of the data sheet.
                        Sox = docE["Sox"] #The name of the data sheet.

                    docW = Document(SoxFormFile) #Template for the Word file
                    docW._body.clear_content()
                    HeaderTable = docW.add_table(rows=1, cols=1)
                    Templatetable = docW.add_table(rows=1, cols=10, style="Table Grid")

                    """
                    FIRST TABLE/FIRST TABLE/FIRST TABLE
                    """
                    HeaderTable.rows[0].cells[0].text = "Sox Form Date: " + str(time.strftime("%d/%m/%Y"))
                    HeaderTable.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(14)
                    HeaderTable.rows[0].cells[0].paragraphs[0].runs[0].bold = True
                    HeaderTable.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    Templatetable.rows[0].cells[0].text = "First Name"
                    Templatetable.rows[0].cells[1].text = "Last Name"
                    Templatetable.rows[0].cells[2].text = "Personal e-mail"
                    Templatetable.rows[0].cells[3].text = "Site Code"
                    Templatetable.rows[0].cells[4].text = "Time Attendance Employee Nr."
                    Templatetable.rows[0].cells[5].text = "PayRoll Employee Nr."
                    Templatetable.rows[0].cells[6].text = "Position"
                    Templatetable.rows[0].cells[7].text = "Department"
                    Templatetable.rows[0].cells[8].text = "Starting / Leaving (Alta / Baja)"
                    Templatetable.rows[0].cells[9].text = "Date"
                    Templatetable.rows[0].cells[0].paragraphs[0].runs[0].bold = True
                    Templatetable.rows[0].cells[1].paragraphs[0].runs[0].bold = True
                    Templatetable.rows[0].cells[2].paragraphs[0].runs[0].bold = True
                    Templatetable.rows[0].cells[3].paragraphs[0].runs[0].bold = True
                    Templatetable.rows[0].cells[4].paragraphs[0].runs[0].bold = True
                    Templatetable.rows[0].cells[5].paragraphs[0].runs[0].bold = True
                    Templatetable.rows[0].cells[6].paragraphs[0].runs[0].bold = True
                    Templatetable.rows[0].cells[7].paragraphs[0].runs[0].bold = True
                    Templatetable.rows[0].cells[8].paragraphs[0].runs[0].bold = True
                    Templatetable.rows[0].cells[9].paragraphs[0].runs[0].bold = True
                    Templatetable.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
                    Templatetable.rows[0].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
                    Templatetable.rows[0].cells[2].paragraphs[0].runs[0].font.size = Pt(11)
                    Templatetable.rows[0].cells[3].paragraphs[0].runs[0].font.size = Pt(11)
                    Templatetable.rows[0].cells[4].paragraphs[0].runs[0].font.size = Pt(11)
                    Templatetable.rows[0].cells[5].paragraphs[0].runs[0].font.size = Pt(11)
                    Templatetable.rows[0].cells[6].paragraphs[0].runs[0].font.size = Pt(11)
                    Templatetable.rows[0].cells[7].paragraphs[0].runs[0].font.size = Pt(11)
                    Templatetable.rows[0].cells[8].paragraphs[0].runs[0].font.size = Pt(11)
                    Templatetable.rows[0].cells[9].paragraphs[0].runs[0].font.size = Pt(11)
                    Templatetable.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    Templatetable.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    Templatetable.rows[0].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    Templatetable.rows[0].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    Templatetable.rows[0].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    Templatetable.rows[0].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    Templatetable.rows[0].cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    Templatetable.rows[0].cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    Templatetable.rows[0].cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    Templatetable.rows[0].cells[9].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    LineDeleted = 1

                    for x in range(2,Sox.max_row+1):

                        FirstName = Sox.cell(row = x, column = 1).value
                        LastName = Sox.cell(row = x, column = 2).value
                        FirstName = replaceSpecialChars(FirstName)
                        LastName = replaceSpecialChars(LastName)
                        PersonalEmail = Sox.cell(row = x, column = 3).value
                        OtherCompany = Sox.cell(row = x, column = 4).value
                        TimeAttendanceNum = Sox.cell(row = x, column = 5).value
                        EmployeeNum = Sox.cell(row = x, column = 6).value
                        Position = Sox.cell(row = x, column = 7).value
                        Dept = Sox.cell(row = x, column = 8).value
                        StartLeave = Sox.cell(row = x, column = 9).value
                        SLDate = Sox.cell(row = x, column = 10).value

                        if StartLeave == "Deleted":
                            if x == Sox.max_row:
                                if optionModify == "1":
                                    NameFile = TableFiles + "\\" + FileYear + "\\" + str(months[int(FileMonth)]) +  "\\" + TMPExcel1 + ".docx"
                                elif optionModify == "2":
                                    NameFile = TableFiles + "\\" + FileYear + "\\" + str(months[int(FileMonth)]) +  "\\" + TMPExcel2 + ".docx"

                                docW.save(NameFile)
                                break
                            else:
                                LineDeleted = LineDeleted + 1
                                
                                continue
                        else:
                            if StartLeave == "Leaving":
                                DateYForm = SLDate.split("/")[2]
                                DateMForm = SLDate.split("/")[1]
                                CorrectFormDate = SLDate.replace("/","-")
                                os.makedirs(FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)], exist_ok=True)
                                os.makedirs(FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated", exist_ok=True)

                                if charCounter(string = FirstName, char = " ") == 1:
                                    if charCounter(string = LastName, char = " ") == 1:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                    elif charCounter(string = LastName, char = " ") == 2:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "-" + CorrectFormDate + ".docx"
                                    elif charCounter(string = LastName, char = " ") == 3:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "_" + LastName.split(" ")[3] + "-" + CorrectFormDate + ".docx"
                                    else:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName + "-" + CorrectFormDate + ".docx"
                                elif charCounter(string = FirstName, char = " ") == 2:
                                    if charCounter(string = LastName, char = " ") == 1:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                    elif charCounter(string = LastName, char = " ") == 2:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "-" + CorrectFormDate + ".docx"
                                    elif charCounter(string = LastName, char = " ") == 3:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "_" + LastName.split(" ")[3] + "-" + CorrectFormDate + ".docx"
                                    else:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName + "-" + CorrectFormDate + ".docx"
                                elif charCounter(string = FirstName, char = " ") == 3:
                                    if charCounter(string = LastName, char = " ") == 1:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                    elif charCounter(string = LastName, char = " ") == 2:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "-" + CorrectFormDate + ".docx"
                                    elif charCounter(string = LastName, char = " ") == 3:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "_" + LastName.split(" ")[3] + "-" + CorrectFormDate + ".docx"
                                    else:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName + "-" + CorrectFormDate + ".docx"

                                else:
                                    if charCounter(string = LastName, char = " ") == 1:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                    else:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName + "_" + LastName + "-" + CorrectFormDate + ".docx"
                                
                                terminatingUsers(FirstName, LastName, EmployeeNum, SLDate, NameTerminatedFile)

                            NumLines = x-LineDeleted
                            if optionModify == "1":
                                NameFile = TableFiles + "\\" + FileYear + "\\" + str(months[int(FileMonth)]) +  "\\" + TMPExcel1 + ".docx"
                            elif optionModify == "2":
                                NameFile = TableFiles + "\\" + FileYear + "\\" + str(months[int(FileMonth)]) +  "\\" + TMPExcel2 + ".docx"

                            Templatetable.add_row()
                            Templatetable.rows[NumLines].cells[0].text = FirstName
                            Templatetable.rows[NumLines].cells[1].text = LastName
                            Templatetable.rows[NumLines].cells[2].text = PersonalEmail
                            Templatetable.rows[NumLines].cells[3].text = OtherCompany
                            Templatetable.rows[NumLines].cells[4].text = TimeAttendanceNum
                            Templatetable.rows[NumLines].cells[5].text = EmployeeNum
                            Templatetable.rows[NumLines].cells[6].text = Position
                            Templatetable.rows[NumLines].cells[7].text = Dept
                            Templatetable.rows[NumLines].cells[8].text = StartLeave
                            Templatetable.rows[NumLines].cells[9].text = SLDate
                            Templatetable.rows[NumLines].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
                            Templatetable.rows[NumLines].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
                            Templatetable.rows[NumLines].cells[2].paragraphs[0].runs[0].font.size = Pt(11)
                            Templatetable.rows[NumLines].cells[3].paragraphs[0].runs[0].font.size = Pt(11)
                            Templatetable.rows[NumLines].cells[4].paragraphs[0].runs[0].font.size = Pt(11)
                            Templatetable.rows[NumLines].cells[5].paragraphs[0].runs[0].font.size = Pt(11)
                            Templatetable.rows[NumLines].cells[6].paragraphs[0].runs[0].font.size = Pt(11)
                            Templatetable.rows[NumLines].cells[7].paragraphs[0].runs[0].font.size = Pt(11)
                            Templatetable.rows[NumLines].cells[8].paragraphs[0].runs[0].font.size = Pt(11)
                            if StartLeave == "Leaving":
                                Templatetable.rows[NumLines].cells[8].paragraphs[0].runs[0].font.color.rgb = RGBColor(200, 0, 0)
                            else:
                                Templatetable.rows[NumLines].cells[8].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 200, 0)
                            Templatetable.rows[NumLines].cells[9].paragraphs[0].runs[0].font.size = Pt(11)
                            Templatetable.rows[NumLines].cells[8].paragraphs[0].runs[0].bold = True
                            Templatetable.rows[NumLines].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            Templatetable.rows[NumLines].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            Templatetable.rows[NumLines].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            Templatetable.rows[NumLines].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            Templatetable.rows[NumLines].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            Templatetable.rows[NumLines].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            Templatetable.rows[NumLines].cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            Templatetable.rows[NumLines].cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            Templatetable.rows[NumLines].cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            Templatetable.rows[NumLines].cells[9].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        if optionModify == "1":
                            NameFile = TableFiles + "\\" + FileYear + "\\" + str(months[int(FileMonth)]) +  "\\" + TMPExcel1 + ".docx"
                        elif optionModify == "2":
                            NameFile = TableFiles + "\\" + FileYear + "\\" + str(months[int(FileMonth)]) +  "\\" + TMPExcel2 + ".docx"

                        docW.save(NameFile)

                    askForSend = boolbox(msg="Do you want to send this file by Email?")
                    if askForSend == True:
                        sendEmail(NameFile=NameFile, FormDayDate=FileDay, ModifyFile = True)
            elif AnswerOption == "2" or AnswerOption == "3":
                SearchUserData = multenterbox(msg="Enter the User First Name and Last Name", title=ProgramTitle, fields=ModifyUserSoxFile)
                SearchFirstName = (SearchUserData[0]).strip()
                SearchLastName = (SearchUserData[1]).strip()
                SearchFirstName = SearchFirstName.capitalize()
                SearchLastName = SearchLastName.capitalize()
                NoneValue = None

                for x in range(2,(Sox.max_row+1)):
                    DataUsers = ["No Data", "No Data", "-", "-", "XXXX", "YYYY", "No Data", "No Data", "No Data", "DD/MM/YYYY"]
                    DataUsers[0] = Sox.cell(row = x, column = 1).value
                    DataUsers[1] = Sox.cell(row = x, column = 2).value
                    DataUsers[2] = Sox.cell(row = x, column = 3).value
                    DataUsers[3] = Sox.cell(row = x, column = 4).value
                    DataUsers[4] = Sox.cell(row = x, column = 5).value
                    DataUsers[5] = Sox.cell(row = x, column = 6).value
                    DataUsers[6] = Sox.cell(row = x, column = 7).value
                    DataUsers[7] = Sox.cell(row = x, column = 8).value
                    DataUsers[8] = Sox.cell(row = x, column = 9).value
                    DataUsers[9] = Sox.cell(row = x, column = 10).value

                    z = 0
                    for y in DataUsers:
                        z += 1
                        print(str(z) + ": " + y)

                    if SearchFirstName == DataUsers[0] and SearchLastName == DataUsers[1]:

                        if AnswerOption == "2":
                            ModifyingSoxValues = multenterbox(msg="I Found this Values. Change the value to modify it", title=ProgramTitle, fields=ModifyDataSoxFileFields, values=DataUsers)

                            FirstName = ModifyingSoxValues[0]
                            LastName = ModifyingSoxValues[1]
                            PersonalEmail = ModifyingSoxValues[2]
                            OtherCompany = ModifyingSoxValues[3]
                            TimeAttendanceNum = ModifyingSoxValues[4]
                            EmployeeNum = ModifyingSoxValues[5]
                            Position = ModifyingSoxValues[6]
                            Dept = ModifyingSoxValues[8]
                            StartLeave = DataUsers[8]
                            SLDate = DataUsers[9]

                            FirstName = FirstName.capitalize()
                            LastName = LastName.capitalize()
                            PersonalEmail = PersonalEmail.capitalize()
                            OtherCompany = OtherCompany.upper()
                            TimeAttendanceNum = TimeAttendanceNum
                            EmployeeNum = EmployeeNum
                            Position = Position.capitalize()

                            if len(FirstName) == 0:
                                FirstName = DataUsers[0]
                            if charCounter(FirstName) == 1:
                                FirstName = FirstName.split(" ")[0].capitalize() + " " +  FirstName.split(" ")[1].capitalize()
                            elif charCounter(FirstName) == 2:
                                FirstName = FirstName.split(" ")[0].capitalize() + " " +  FirstName.split(" ")[1].capitalize() + " " +  FirstName.split(" ")[2].capitalize()

                            if len(LastName) == 0:
                                LastName = DataUsers[1]
                            if charCounter(LastName) == 1:
                                LastName = LastName.split(" ")[0].capitalize() + " " +  LastName.split(" ")[1].capitalize()
                            elif charCounter(LastName) == 2:
                                LastName = LastName.split(" ")[0].capitalize() + " " +  LastName.split(" ")[1].capitalize() + " " +  LastName.split(" ")[2].capitalize()
                            
                            FirstName = replaceSpecialChars(FirstName)
                            LastName = replaceSpecialChars(LastName)

                            if len(PersonalEmail) == 0 or PersonalEmail == DefaultSoxData[2]:
                                PersonalEmail = DataUsers[2]

                            if len(OtherCompany) == 0:
                                OtherCompany = DataUsers[3]
                            if OtherCompany != DataUsers[3]:
                                OtherCompany = OtherCompany.upper()
                            else:
                                OtherCompany = DataUsers[3]

                            if len(TimeAttendanceNum) == 0 or TimeAttendanceNum == "XXXX":
                                TimeAttendanceNum = DataUsers[4]

                            if len(EmployeeNum) == 0 or EmployeeNum == "YYYY":
                                EmployeeNum = DataUsers[5]

                            if len(Position) == 0:
                                Position = DataUsers[6]
                            if charCounter(Position) == 1:
                                Position = Position.split(" ")[0].capitalize() + " " +  Position.split(" ")[1].capitalize()
                            elif charCounter(Position) == 2:
                                Position = Position.split(" ")[0].capitalize() + " " +  Position.split(" ")[1].capitalize() + " " +  Position.split(" ")[2].capitalize()

                            if len(StartLeave) == 0:
                                StartLeave = DataUsers[7]
                            elif StartLeave.capitalize() == "Entry":
                                StartLeave = "Starting"
                            elif StartLeave != "Starting" or StartLeave != "Leaving":
                                StartLeave = ((buttonbox(msg="Is the user \"Starting or Leaving\"", title=ProgramTitle, choices=["[1] Starting", "[2] Leaving", "[3] Cancel"])).split("] ")[0]).replace("[","")
                                if StartLeave == "1":
                                    StartLeave = "Starting"
                                elif StartLeave == "2":
                                    StartLeave = "Leaving"
                                else:
                                    msgbox(msg="Starting will be used as default Starting/Leaving data")
                                    StartLeave = "Starting"

                            if len(SLDate) == 0 or SLDate =="DD/MM/YYYY":
                                SLDate = DataUsers[8]

                            Sox.cell(row=x,column=1, value=FirstName)
                            Sox.cell(row=x,column=2, value=LastName)
                            Sox.cell(row=x,column=3, value=PersonalEmail)
                            Sox.cell(row=x,column=4, value=OtherCompany)
                            Sox.cell(row=x,column=5, value=TimeAttendanceNum)
                            Sox.cell(row=x,column=6, value=EmployeeNum)
                            Sox.cell(row=x,column=7, value=Position)
                            Sox.cell(row=x,column=8, value=Dept)
                            Sox.cell(row=x,column=9, value=StartLeave)
                            Sox.cell(row=x,column=10, value=SLDate)

                            if optionModify == "1":
                                docE.save(TMPExcel)
                                docE = load_workbook(TMPExcel) #File where we have the data saved.
                                Data = docE["Users"] #The name of the data sheet.
                                Sox = docE["Sox"] #The name of the data sheet.
                            elif optionModify == "2":
                                docE.save(TMPExcelModified) #File where we have the data saved.
                                docE = load_workbook(TMPExcelModified) #File where we have the data saved.
                                Data = docE["Users"] #The name of the data sheet.
                                Sox = docE["Sox"] #The name of the data sheet.

                            docW = Document(SoxFormFile)
                            docW._body.clear_content()
                            HeaderTable = docW.add_table(rows=1, cols=1)
                            Templatetable = docW.add_table(rows=1, cols=10, style="Table Grid")

                            """
                            FIRST TABLE/FIRST TABLE/FIRST TABLE
                            """
                            HeaderTable.rows[0].cells[0].text = "Sox Form Date: " + str(time.strftime("%d/%m/%Y"))
                            HeaderTable.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(14)
                            HeaderTable.rows[0].cells[0].paragraphs[0].runs[0].bold = True
                            HeaderTable.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                            Templatetable.rows[0].cells[0].text = "First Name"
                            Templatetable.rows[0].cells[1].text = "Last Name"
                            Templatetable.rows[0].cells[2].text = "Personal e-mail"
                            Templatetable.rows[0].cells[3].text = "Site Code"
                            Templatetable.rows[0].cells[4].text = "Time Attendance Employee Nr."
                            Templatetable.rows[0].cells[5].text = "PayRoll Employee Nr."
                            Templatetable.rows[0].cells[6].text = "Position"
                            Templatetable.rows[0].cells[7].text = "Department"
                            Templatetable.rows[0].cells[8].text = "Starting / Leaving (Alta / Baja)"
                            Templatetable.rows[0].cells[9].text = "Date"
                            Templatetable.rows[0].cells[0].paragraphs[0].runs[0].bold = True
                            Templatetable.rows[0].cells[1].paragraphs[0].runs[0].bold = True
                            Templatetable.rows[0].cells[2].paragraphs[0].runs[0].bold = True
                            Templatetable.rows[0].cells[3].paragraphs[0].runs[0].bold = True
                            Templatetable.rows[0].cells[4].paragraphs[0].runs[0].bold = True
                            Templatetable.rows[0].cells[5].paragraphs[0].runs[0].bold = True
                            Templatetable.rows[0].cells[6].paragraphs[0].runs[0].bold = True
                            Templatetable.rows[0].cells[7].paragraphs[0].runs[0].bold = True
                            Templatetable.rows[0].cells[8].paragraphs[0].runs[0].bold = True
                            Templatetable.rows[0].cells[9].paragraphs[0].runs[0].bold = True
                            Templatetable.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
                            Templatetable.rows[0].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
                            Templatetable.rows[0].cells[2].paragraphs[0].runs[0].font.size = Pt(11)
                            Templatetable.rows[0].cells[3].paragraphs[0].runs[0].font.size = Pt(11)
                            Templatetable.rows[0].cells[4].paragraphs[0].runs[0].font.size = Pt(11)
                            Templatetable.rows[0].cells[5].paragraphs[0].runs[0].font.size = Pt(11)
                            Templatetable.rows[0].cells[6].paragraphs[0].runs[0].font.size = Pt(11)
                            Templatetable.rows[0].cells[7].paragraphs[0].runs[0].font.size = Pt(11)
                            Templatetable.rows[0].cells[8].paragraphs[0].runs[0].font.size = Pt(11)
                            Templatetable.rows[0].cells[9].paragraphs[0].runs[0].font.size = Pt(11)
                            Templatetable.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            Templatetable.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            Templatetable.rows[0].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            Templatetable.rows[0].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            Templatetable.rows[0].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            Templatetable.rows[0].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            Templatetable.rows[0].cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            Templatetable.rows[0].cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            Templatetable.rows[0].cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            Templatetable.rows[0].cells[9].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            LineDeleted = 1

                            for x in range(2,Sox.max_row+1):

                                FirstName = Sox.cell(row = x, column = 1).value
                                LastName = Sox.cell(row = x, column = 2).value
                                FirstName = replaceSpecialChars(FirstName)
                                LastName = replaceSpecialChars(LastName)
                                PersonalEmail = Sox.cell(row = x, column = 3).value
                                OtherCompany = Sox.cell(row = x, column = 4).value
                                TimeAttendanceNum = Sox.cell(row = x, column = 5).value
                                EmployeeNum = Sox.cell(row = x, column = 6).value
                                Position = Sox.cell(row = x, column = 7).value
                                Dept = Sox.cell(row = x, column = 8).value
                                StartLeave = Sox.cell(row = x, column = 9).value
                                SLDate = Sox.cell(row = x, column = 10).value
                                if StartLeave == "Deleted":
                                    if x == Sox.max_row:
                                        if optionModify == "1":
                                            NameFile = TableFiles + "\\" + FileYear + "\\" + str(months[int(FileMonth)]) +  "\\" + TMPExcel1 + ".docx"
                                        elif optionModify == "2":
                                            NameFile = TableFiles + "\\" + FileYear + "\\" + str(months[int(FileMonth)]) +  "\\" + TMPExcel2 + ".docx"

                                        docW.save(NameFile)
                                        break
                                    else:
                                        LineDeleted = LineDeleted + 1
                                        continue
                                else:
                                    if StartLeave == "Leaving":
                                        DateYForm = SLDate.split("/")[2]
                                        DateMForm = SLDate.split("/")[1]
                                        CorrectFormDate = SLDate.replace("/","-")
                                        os.makedirs(FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)], exist_ok=True)
                                        os.makedirs(FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated", exist_ok=True)

                                        if charCounter(string = FirstName, char = " ") == 1:
                                            if charCounter(string = LastName, char = " ") == 1:
                                                NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                            elif charCounter(string = LastName, char = " ") == 2:
                                                NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "-" + CorrectFormDate + ".docx"
                                            elif charCounter(string = LastName, char = " ") == 3:
                                                NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "_" + LastName.split(" ")[3] + "-" + CorrectFormDate + ".docx"
                                            else:
                                                NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName + "-" + CorrectFormDate + ".docx"
                                        elif charCounter(string = FirstName, char = " ") == 2:
                                            if charCounter(string = LastName, char = " ") == 1:
                                                NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                            elif charCounter(string = LastName, char = " ") == 2:
                                                NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "-" + CorrectFormDate + ".docx"
                                            elif charCounter(string = LastName, char = " ") == 3:
                                                NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "_" + LastName.split(" ")[3] + "-" + CorrectFormDate + ".docx"
                                            else:
                                                NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName + "-" + CorrectFormDate + ".docx"
                                        elif charCounter(string = FirstName, char = " ") == 3:
                                            if charCounter(string = LastName, char = " ") == 1:
                                                NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                            elif charCounter(string = LastName, char = " ") == 2:
                                                NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "-" + CorrectFormDate + ".docx"
                                            elif charCounter(string = LastName, char = " ") == 3:
                                                NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "_" + LastName.split(" ")[3] + "-" + CorrectFormDate + ".docx"
                                            else:
                                                NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName + "-" + CorrectFormDate + ".docx"

                                        else:
                                            if charCounter(string = LastName, char = " ") == 1:
                                                NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                            else:
                                                NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName + "_" + LastName + "-" + CorrectFormDate + ".docx"
                                                                
                                        terminatingUsers(FirstName, LastName, EmployeeNum, SLDate, NameTerminatedFile)

                                    NumLines = x-LineDeleted
                                    if optionModify == "1":
                                        NameFile = TableFiles + "\\" + FileYear + "\\" + str(months[int(FileMonth)]) +  "\\" + TMPExcel1 + ".docx"
                                    elif optionModify == "2":
                                        NameFile = TableFiles + "\\" + FileYear + "\\" + str(months[int(FileMonth)]) +  "\\" + TMPExcel2 + ".docx"

                                    Templatetable.add_row()
                                    Templatetable.rows[NumLines].cells[0].text = FirstName
                                    Templatetable.rows[NumLines].cells[1].text = LastName
                                    Templatetable.rows[NumLines].cells[2].text = PersonalEmail
                                    Templatetable.rows[NumLines].cells[3].text = OtherCompany
                                    Templatetable.rows[NumLines].cells[4].text = TimeAttendanceNum
                                    Templatetable.rows[NumLines].cells[5].text = EmployeeNum
                                    Templatetable.rows[NumLines].cells[6].text = Position
                                    Templatetable.rows[NumLines].cells[7].text = Dept
                                    Templatetable.rows[NumLines].cells[8].text = StartLeave
                                    Templatetable.rows[NumLines].cells[9].text = SLDate
                                    Templatetable.rows[NumLines].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
                                    Templatetable.rows[NumLines].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
                                    Templatetable.rows[NumLines].cells[2].paragraphs[0].runs[0].font.size = Pt(11)
                                    Templatetable.rows[NumLines].cells[3].paragraphs[0].runs[0].font.size = Pt(11)
                                    Templatetable.rows[NumLines].cells[4].paragraphs[0].runs[0].font.size = Pt(11)
                                    Templatetable.rows[NumLines].cells[5].paragraphs[0].runs[0].font.size = Pt(11)
                                    Templatetable.rows[NumLines].cells[6].paragraphs[0].runs[0].font.size = Pt(11)
                                    Templatetable.rows[NumLines].cells[7].paragraphs[0].runs[0].font.size = Pt(11)
                                    Templatetable.rows[NumLines].cells[8].paragraphs[0].runs[0].font.size = Pt(11)
                                    if StartLeave == "Leaving":
                                        Templatetable.rows[NumLines].cells[8].paragraphs[0].runs[0].font.color.rgb = RGBColor(200, 0, 0)
                                    else:
                                        Templatetable.rows[NumLines].cells[8].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 200, 0)
                                    Templatetable.rows[NumLines].cells[9].paragraphs[0].runs[0].font.size = Pt(11)
                                    Templatetable.rows[NumLines].cells[8].paragraphs[0].runs[0].bold = True
                                    Templatetable.rows[NumLines].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    Templatetable.rows[NumLines].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    Templatetable.rows[NumLines].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    Templatetable.rows[NumLines].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    Templatetable.rows[NumLines].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    Templatetable.rows[NumLines].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    Templatetable.rows[NumLines].cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    Templatetable.rows[NumLines].cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    Templatetable.rows[NumLines].cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    Templatetable.rows[NumLines].cells[9].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            else:

                                if optionModify == "1":
                                    NameFile = TableFiles + "\\" + FileYear + "\\" + str(months[int(FileMonth)]) +  "\\" + TMPExcel1 + ".docx"
                                elif optionModify == "2":
                                    NameFile = TableFiles + "\\" + FileYear + "\\" + str(months[int(FileMonth)]) +  "\\" + TMPExcel2 + ".docx"
                                docW.save(NameFile)

                                askForSend = boolbox(msg="Do you want Send this file by Email?")

                                if askForSend == True:
                                    sendEmail(NameFile=NameFile, FormDayDate=FileDay, ModifyFile = True)
                        elif AnswerOption == "3":
                            AskingToDelete = boolbox(title=ProgramTitle, msg="I Found this data\n\n" + "First Name: " + DataUsers[0] + "\n" + "Last Name: " + DataUsers[1] + "\n" + "Personal Email: " + DataUsers[2] + "\n" + "Site Code: " + DataUsers[3] + "\n" + "Time Attendance Number: " + DataUsers[4] + "\n" + "Employee Number: " + DataUsers[5] + "\n" + "Position: " + DataUsers[6] + "\n" + "Department: " + DataUsers[7] + "\n" + "Starting/Leaving: " + DataUsers[8] + "\n" + "Date: " + DataUsers[9] + "\n\n\n" + "Would you like to delete " + DataUsers[0] + " " + DataUsers[1] + " from Sox File?")

                            if AskingToDelete == False:
                                break
                            else:
                                Sox.cell(row=x,column=1, value=DataUsers[0])
                                Sox.cell(row=x,column=2, value=DataUsers[1])
                                Sox.cell(row=x,column=3, value="Deleted " + time.strftime("%d/%m/%Y"))
                                Sox.cell(row=x,column=4, value="Deleted " + time.strftime("%d/%m/%Y"))
                                Sox.cell(row=x,column=5, value="Deleted " + time.strftime("%d/%m/%Y"))
                                Sox.cell(row=x,column=6, value="Deleted " + time.strftime("%d/%m/%Y"))
                                Sox.cell(row=x,column=7, value="Deleted " + time.strftime("%d/%m/%Y"))
                                Sox.cell(row=x,column=8, value="Deleted " + time.strftime("%d/%m/%Y"))
                                Sox.cell(row=x,column=9, value="Deleted")
                                Sox.cell(row=x,column=10, value="Deleted " + time.strftime("%d/%m/%Y"))

                                if optionModify == "1":
                                    docE.save(TMPExcel)
                                    docE = load_workbook(TMPExcel) #File where we have the data saved.
                                    Data = docE["Users"] #The name of the data sheet.
                                    Sox = docE["Sox"] #The name of the data sheet.
                                elif optionModify == "2":
                                    docE.save(TMPExcelModified) #File where we have the data saved.
                                    docE = load_workbook(TMPExcelModified) #File where we have the data saved.
                                    Data = docE["Users"] #The name of the data sheet.
                                    Sox = docE["Sox"] #The name of the data sheet.

                                docW = Document(SoxFormFile)
                                docW._body.clear_content()
                                HeaderTable = docW.add_table(rows=1, cols=1)
                                Templatetable = docW.add_table(rows=1, cols=10, style="Table Grid")

                                """
                                FIRST TABLE/FIRST TABLE/FIRST TABLE
                                """
                                HeaderTable.rows[0].cells[0].text = "Sox Form Date: " + str(time.strftime("%d/%m/%Y"))
                                HeaderTable.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(14)
                                HeaderTable.rows[0].cells[0].paragraphs[0].runs[0].bold = True
                                HeaderTable.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                                Templatetable.rows[0].cells[0].text = "First Name"
                                Templatetable.rows[0].cells[1].text = "Last Name"
                                Templatetable.rows[0].cells[2].text = "Personal e-mail"
                                Templatetable.rows[0].cells[3].text = "Site Code"
                                Templatetable.rows[0].cells[4].text = "Time Attendance Employee Nr."
                                Templatetable.rows[0].cells[5].text = "PayRoll Employee Nr."
                                Templatetable.rows[0].cells[6].text = "Position"
                                Templatetable.rows[0].cells[7].text = "Department"
                                Templatetable.rows[0].cells[8].text = "Starting / Leaving (Alta / Baja)"
                                Templatetable.rows[0].cells[9].text = "Date"
                                Templatetable.rows[0].cells[0].paragraphs[0].runs[0].bold = True
                                Templatetable.rows[0].cells[1].paragraphs[0].runs[0].bold = True
                                Templatetable.rows[0].cells[2].paragraphs[0].runs[0].bold = True
                                Templatetable.rows[0].cells[3].paragraphs[0].runs[0].bold = True
                                Templatetable.rows[0].cells[4].paragraphs[0].runs[0].bold = True
                                Templatetable.rows[0].cells[5].paragraphs[0].runs[0].bold = True
                                Templatetable.rows[0].cells[6].paragraphs[0].runs[0].bold = True
                                Templatetable.rows[0].cells[7].paragraphs[0].runs[0].bold = True
                                Templatetable.rows[0].cells[8].paragraphs[0].runs[0].bold = True
                                Templatetable.rows[0].cells[9].paragraphs[0].runs[0].bold = True
                                Templatetable.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
                                Templatetable.rows[0].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
                                Templatetable.rows[0].cells[2].paragraphs[0].runs[0].font.size = Pt(11)
                                Templatetable.rows[0].cells[3].paragraphs[0].runs[0].font.size = Pt(11)
                                Templatetable.rows[0].cells[4].paragraphs[0].runs[0].font.size = Pt(11)
                                Templatetable.rows[0].cells[5].paragraphs[0].runs[0].font.size = Pt(11)
                                Templatetable.rows[0].cells[6].paragraphs[0].runs[0].font.size = Pt(11)
                                Templatetable.rows[0].cells[7].paragraphs[0].runs[0].font.size = Pt(11)
                                Templatetable.rows[0].cells[8].paragraphs[0].runs[0].font.size = Pt(11)
                                Templatetable.rows[0].cells[9].paragraphs[0].runs[0].font.size = Pt(11)
                                Templatetable.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                Templatetable.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                Templatetable.rows[0].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                Templatetable.rows[0].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                Templatetable.rows[0].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                Templatetable.rows[0].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                Templatetable.rows[0].cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                Templatetable.rows[0].cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                Templatetable.rows[0].cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                Templatetable.rows[0].cells[9].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                LineDeleted = 1

                                for x in range(2,Sox.max_row+1):

                                    FirstName = Sox.cell(row = x, column = 1).value
                                    LastName = Sox.cell(row = x, column = 2).value
                                    FirstName = replaceSpecialChars(FirstName)
                                    LastName = replaceSpecialChars(LastName)
                                    PersonalEmail = Sox.cell(row = x, column = 3).value
                                    OtherCompany = Sox.cell(row = x, column = 4).value
                                    TimeAttendanceNum = Sox.cell(row = x, column = 5).value
                                    EmployeeNum = Sox.cell(row = x, column = 6).value
                                    Position = Sox.cell(row = x, column = 7).value
                                    Dept = Sox.cell(row = x, column = 8).value
                                    StartLeave = Sox.cell(row = x, column = 9).value
                                    SLDate = Sox.cell(row = x, column = 10).value
                                    if StartLeave == "Deleted":
                                        if x == Sox.max_row:
                                            if optionModify == "1":
                                                NameFile = TableFiles + "\\" + FileYear + "\\" + str(months[int(FileMonth)]) +  "\\" + TMPExcel1 + ".docx"
                                            elif optionModify == "2":
                                                NameFile = TableFiles + "\\" + FileYear + "\\" + str(months[int(FileMonth)]) +  "\\" + TMPExcel2 + ".docx"

                                            docW.save(NameFile)
                                            break
                                        else:
                                            LineDeleted = LineDeleted + 1
                                            continue
                                    else:
                                        NumLines = x-LineDeleted

                                        Templatetable.add_row()
                                        Templatetable.rows[NumLines].cells[0].text = FirstName
                                        Templatetable.rows[NumLines].cells[1].text = LastName
                                        Templatetable.rows[NumLines].cells[2].text = PersonalEmail
                                        Templatetable.rows[NumLines].cells[3].text = OtherCompany
                                        Templatetable.rows[NumLines].cells[4].text = TimeAttendanceNum
                                        Templatetable.rows[NumLines].cells[5].text = EmployeeNum
                                        Templatetable.rows[NumLines].cells[6].text = Position
                                        Templatetable.rows[NumLines].cells[7].text = Dept
                                        Templatetable.rows[NumLines].cells[8].text = StartLeave
                                        Templatetable.rows[NumLines].cells[9].text = SLDate
                                        Templatetable.rows[NumLines].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
                                        Templatetable.rows[NumLines].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
                                        Templatetable.rows[NumLines].cells[2].paragraphs[0].runs[0].font.size = Pt(11)
                                        Templatetable.rows[NumLines].cells[3].paragraphs[0].runs[0].font.size = Pt(11)
                                        Templatetable.rows[NumLines].cells[4].paragraphs[0].runs[0].font.size = Pt(11)
                                        Templatetable.rows[NumLines].cells[5].paragraphs[0].runs[0].font.size = Pt(11)
                                        Templatetable.rows[NumLines].cells[6].paragraphs[0].runs[0].font.size = Pt(11)
                                        Templatetable.rows[NumLines].cells[7].paragraphs[0].runs[0].font.size = Pt(11)
                                        Templatetable.rows[NumLines].cells[8].paragraphs[0].runs[0].font.size = Pt(11)
                                        if StartLeave == "Leaving":
                                            Templatetable.rows[NumLines].cells[8].paragraphs[0].runs[0].font.color.rgb = RGBColor(200, 0, 0)
                                        else:
                                            Templatetable.rows[NumLines].cells[8].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 200, 0)
                                        Templatetable.rows[NumLines].cells[9].paragraphs[0].runs[0].font.size = Pt(11)
                                        Templatetable.rows[NumLines].cells[8].paragraphs[0].runs[0].bold = True
                                        Templatetable.rows[NumLines].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        Templatetable.rows[NumLines].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        Templatetable.rows[NumLines].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        Templatetable.rows[NumLines].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        Templatetable.rows[NumLines].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        Templatetable.rows[NumLines].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        Templatetable.rows[NumLines].cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        Templatetable.rows[NumLines].cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        Templatetable.rows[NumLines].cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        Templatetable.rows[NumLines].cells[9].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                else:
                                    if optionModify == "1":
                                        NameFile = TableFiles + "\\" + FileYear + "\\" + str(months[int(FileMonth)]) +  "\\" + TMPExcel1 + ".docx"
                                    elif optionModify == "2":
                                        NameFile = TableFiles + "\\" + FileYear + "\\" + str(months[int(FileMonth)]) +  "\\" + TMPExcel2 + ".docx"
                                    
                                    docW.save(NameFile)

                                    askForSend = boolbox(msg="Do you want Send this file by Email?")
                                    if askForSend == True:
                                        sendEmail(NameFile=NameFile, FormDayDate=FileDay, ModifyFile = True)
            elif AnswerOption == "4":

                docW = Document(SoxFormFile)
                docW._body.clear_content()
                HeaderTable = docW.add_table(rows=1, cols=1)
                Templatetable = docW.add_table(rows=1, cols=10, style="Table Grid")

                """
                FIRST TABLE/FIRST TABLE/FIRST TABLE
                """
                HeaderTable.rows[0].cells[0].text = "Sox Form Date: " + str(time.strftime("%d/%m/%Y"))
                HeaderTable.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(14)
                HeaderTable.rows[0].cells[0].paragraphs[0].runs[0].bold = True
                HeaderTable.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                Templatetable.rows[0].cells[0].text = "First Name"
                Templatetable.rows[0].cells[1].text = "Last Name"
                Templatetable.rows[0].cells[2].text = "Personal e-mail"
                Templatetable.rows[0].cells[3].text = "Site Code"
                Templatetable.rows[0].cells[4].text = "Time Attendance Employee Nr."
                Templatetable.rows[0].cells[5].text = "PayRoll Employee Nr."
                Templatetable.rows[0].cells[6].text = "Position"
                Templatetable.rows[0].cells[7].text = "Department"
                Templatetable.rows[0].cells[8].text = "Starting / Leaving (Alta / Baja)"
                Templatetable.rows[0].cells[9].text = "Date"
                Templatetable.rows[0].cells[0].paragraphs[0].runs[0].bold = True
                Templatetable.rows[0].cells[1].paragraphs[0].runs[0].bold = True
                Templatetable.rows[0].cells[2].paragraphs[0].runs[0].bold = True
                Templatetable.rows[0].cells[3].paragraphs[0].runs[0].bold = True
                Templatetable.rows[0].cells[4].paragraphs[0].runs[0].bold = True
                Templatetable.rows[0].cells[5].paragraphs[0].runs[0].bold = True
                Templatetable.rows[0].cells[6].paragraphs[0].runs[0].bold = True
                Templatetable.rows[0].cells[7].paragraphs[0].runs[0].bold = True
                Templatetable.rows[0].cells[8].paragraphs[0].runs[0].bold = True
                Templatetable.rows[0].cells[9].paragraphs[0].runs[0].bold = True
                Templatetable.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
                Templatetable.rows[0].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
                Templatetable.rows[0].cells[2].paragraphs[0].runs[0].font.size = Pt(11)
                Templatetable.rows[0].cells[3].paragraphs[0].runs[0].font.size = Pt(11)
                Templatetable.rows[0].cells[4].paragraphs[0].runs[0].font.size = Pt(11)
                Templatetable.rows[0].cells[5].paragraphs[0].runs[0].font.size = Pt(11)
                Templatetable.rows[0].cells[6].paragraphs[0].runs[0].font.size = Pt(11)
                Templatetable.rows[0].cells[7].paragraphs[0].runs[0].font.size = Pt(11)
                Templatetable.rows[0].cells[8].paragraphs[0].runs[0].font.size = Pt(11)
                Templatetable.rows[0].cells[9].paragraphs[0].runs[0].font.size = Pt(11)
                Templatetable.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                Templatetable.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                Templatetable.rows[0].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                Templatetable.rows[0].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                Templatetable.rows[0].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                Templatetable.rows[0].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                Templatetable.rows[0].cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                Templatetable.rows[0].cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                Templatetable.rows[0].cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                Templatetable.rows[0].cells[9].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                LineDeleted = 1

                for x in range(2,Sox.max_row+1):

                    FirstName = Sox.cell(row = x, column = 1).value
                    LastName = Sox.cell(row = x, column = 2).value
                    FirstName = replaceSpecialChars(FirstName)
                    LastName = replaceSpecialChars(LastName)
                    PersonalEmail = Sox.cell(row = x, column = 3).value
                    OtherCompany = Sox.cell(row = x, column = 4).value
                    TimeAttendanceNum = Sox.cell(row = x, column = 5).value
                    EmployeeNum = Sox.cell(row = x, column = 6).value
                    Position = Sox.cell(row = x, column = 7).value
                    Dept = Sox.cell(row = x, column = 8).value
                    StartLeave = Sox.cell(row = x, column = 9).value
                    SLDate = Sox.cell(row = x, column = 10).value
                    if StartLeave == "Deleted":
                        if x == Sox.max_row:
                            if optionModify == "1":
                                NameFile = TableFiles + "\\" + FileYear + "\\" + str(months[int(FileMonth)]) +  "\\" + TMPExcel1 + ".docx"
                            elif optionModify == "2":
                                NameFile = TableFiles + "\\" + FileYear + "\\" + str(months[int(FileMonth)]) +  "\\" + TMPExcel2 + ".docx"

                            docW.save(NameFile)
                            break
                        else:
                            LineDeleted = LineDeleted + 1
                            continue
                    else:
                        if StartLeave == "Leaving":
                            DateYForm = SLDate.split("/")[2]
                            DateMForm = SLDate.split("/")[1]
                            CorrectFormDate = SLDate.replace("/","-")
                            os.makedirs(FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)], exist_ok=True)
                            os.makedirs(FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated", exist_ok=True)

                            if charCounter(string = FirstName, char = " ") == 1:
                                if charCounter(string = LastName, char = " ") == 1:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                elif charCounter(string = LastName, char = " ") == 2:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "-" + CorrectFormDate + ".docx"
                                elif charCounter(string = LastName, char = " ") == 3:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "_" + LastName.split(" ")[3] + "-" + CorrectFormDate + ".docx"
                                else:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName + "-" + CorrectFormDate + ".docx"
                            elif charCounter(string = FirstName, char = " ") == 2:
                                if charCounter(string = LastName, char = " ") == 1:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                elif charCounter(string = LastName, char = " ") == 2:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "-" + CorrectFormDate + ".docx"
                                elif charCounter(string = LastName, char = " ") == 3:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "_" + LastName.split(" ")[3] + "-" + CorrectFormDate + ".docx"
                                else:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName + "-" + CorrectFormDate + ".docx"
                            elif charCounter(string = FirstName, char = " ") == 3:
                                if charCounter(string = LastName, char = " ") == 1:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                elif charCounter(string = LastName, char = " ") == 2:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "-" + CorrectFormDate + ".docx"
                                elif charCounter(string = LastName, char = " ") == 3:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "_" + LastName.split(" ")[3] + "-" + CorrectFormDate + ".docx"
                                else:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName + "-" + CorrectFormDate + ".docx"

                            else:
                                if charCounter(string = LastName, char = " ") == 1:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                else:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName + "_" + LastName + "-" + CorrectFormDate + ".docx"
                                                    
                            terminatingUsers(FirstName, LastName, EmployeeNum, SLDate, NameTerminatedFile)
                            
                        NumLines = x-LineDeleted

                        Templatetable.add_row()
                        Templatetable.rows[NumLines].cells[0].text = FirstName
                        Templatetable.rows[NumLines].cells[1].text = LastName
                        Templatetable.rows[NumLines].cells[2].text = PersonalEmail
                        Templatetable.rows[NumLines].cells[3].text = OtherCompany
                        Templatetable.rows[NumLines].cells[4].text = TimeAttendanceNum
                        Templatetable.rows[NumLines].cells[5].text = EmployeeNum
                        Templatetable.rows[NumLines].cells[6].text = Position
                        Templatetable.rows[NumLines].cells[7].text = Dept
                        Templatetable.rows[NumLines].cells[8].text = StartLeave
                        Templatetable.rows[NumLines].cells[9].text = SLDate
                        Templatetable.rows[NumLines].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
                        Templatetable.rows[NumLines].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
                        Templatetable.rows[NumLines].cells[2].paragraphs[0].runs[0].font.size = Pt(11)
                        Templatetable.rows[NumLines].cells[3].paragraphs[0].runs[0].font.size = Pt(11)
                        Templatetable.rows[NumLines].cells[4].paragraphs[0].runs[0].font.size = Pt(11)
                        Templatetable.rows[NumLines].cells[5].paragraphs[0].runs[0].font.size = Pt(11)
                        Templatetable.rows[NumLines].cells[6].paragraphs[0].runs[0].font.size = Pt(11)
                        Templatetable.rows[NumLines].cells[7].paragraphs[0].runs[0].font.size = Pt(11)
                        Templatetable.rows[NumLines].cells[8].paragraphs[0].runs[0].font.size = Pt(11)
                        if StartLeave == "Leaving":
                            Templatetable.rows[NumLines].cells[8].paragraphs[0].runs[0].font.color.rgb = RGBColor(200, 0, 0)
                        else:
                            Templatetable.rows[NumLines].cells[8].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 200, 0)
                        Templatetable.rows[NumLines].cells[9].paragraphs[0].runs[0].font.size = Pt(11)
                        Templatetable.rows[NumLines].cells[8].paragraphs[0].runs[0].bold = True
                        Templatetable.rows[NumLines].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable.rows[NumLines].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable.rows[NumLines].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable.rows[NumLines].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable.rows[NumLines].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable.rows[NumLines].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable.rows[NumLines].cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable.rows[NumLines].cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable.rows[NumLines].cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable.rows[NumLines].cells[9].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    if optionModify == "1":
                        NameFile = TableFiles + "\\" + FileYear + "\\" + str(months[int(FileMonth)]) +  "\\" + TMPExcel1 + ".docx"
                    elif optionModify == "2":
                        NameFile = TableFiles + "\\" + FileYear + "\\" + str(months[int(FileMonth)]) +  "\\" + TMPExcel2 + ".docx"

                docW.save(NameFile) # Last time, this line it was inside the else from above.
                askForSend = boolbox(msg="Do you want Send this file by Email?")
                if askForSend == True:
                    if optionModify == "1":
                        sendEmail(NameFile=NameFile, FormDayDate=FileDay, ModifyFile = False)
                    elif optionModify == "2":
                        sendEmail(NameFile=NameFile, FormDayDate=FileDay, ModifyFile = True)
            elif AnswerOption == str(5):
                break
        elif MainMenuChoice == "3":
            DateSoxFile = multenterbox(msg="Enter Data for Sox Form File", title=ProgramTitle, fields=ModifySoxFileFields, values=DefaultValuesDataSoxFile)

            FileDay = DateSoxFile[0]
            FileMonth = DateSoxFile[1]
            FileYear = DateSoxFile[2]

            if FileDay[:1] == "0":
                FileDay = FileDay[:1]

            if FileMonth[:1] == "0":
                FileMonth = FileMonth[:1]

            TMPSoxExcelName = "Company Sox " + str(months[int(FileMonth)]) + " - " + FileYear + " - " + FileDay
            TMPSoxExcelName2 = "Company Modified Sox " + str(months[int(FileMonth)]) + " - " + FileYear + " - " + FileDay
            TMPSoxExcel = TableFiles + "\\TMP" + "\\" + TMPSoxExcelName + ".xlsx"
            TMPSoxExcel2 = TableFiles + "\\TMP" + "\\" + TMPSoxExcelName2 + ".xlsx"

            TMPAccessExcelName = "Company_" + FileDay + "-" + FileMonth + "-" + FileYear + ".xlsx"
            TMPAccessExcel = FormFiles + "\\TMP" + "\\" + TMPAccessExcelName
            TMPAccessExcelName2 = "Company_Modified_" + FileDay + "-" + FileMonth + "-" + FileYear + ".xlsx"
            TMPAccessExcel2 = FormFiles + "\\TMP" + "\\" + TMPAccessExcelName2
            DuplicateFileChoices = ["[1] " + TMPAccessExcelName, "[2] " + TMPAccessExcelName2]
            AccessOptions = "0"
            optionModify = "1"

            if os.path.exists(TMPSoxExcel) == True and os.path.exists(TMPSoxExcel2) == True:
                if os.path.exists(TMPAccessExcel2) == True:
                    docE = load_workbook(TMPAccessExcel2)
                    Data = docE["Users"]
                    Sox = docE["Sox"]
                else:
                    optionModify = ((buttonbox(msg=DuplicateFileMsg, title=ProgramTitle, choices=DuplicateFileChoices)).split("] ")[0]).replace("[","")

                    if optionModify == "1":
                        ModifyExcel = 0
                        docE = load_workbook(TMPAccessExcel)
                        AccessFileName = "TMPAccessExcel"
                        Data = docE["Users"]
                        Sox = docE["Sox"]
                    elif optionModify == "2":
                        ModifyExcel = 1
                        shutil.copy2(TMPAccessExcel, TMPAccessExcel2)
                        docE = load_workbook(TMPAccessExcel2)
                        AccessFileName = "TMPAccessExcel2"
                        Data = docE["Users"]
                        Sox = docE["Sox"]
                    else:
                        ModifyExcel = 0
                        docE = load_workbook(TMPAccessExcel)
                        Data = docE["Users"]
                        Sox = docE["Sox"]
            elif os.path.exists(TMPSoxExcel) == True and os.path.exists(TMPSoxExcel2) == False:
                if os.path.exists(TMPAccessExcel) == True:
                    docE = load_workbook(TMPAccessExcel)
                    Data = docE["Users"]
                    Sox = docE["Sox"]
                else:
                    shutil.copy2(TMPSoxExcel, TMPAccessExcel)
                    docE = load_workbook(TMPAccessExcel)
                    Data = docE["Users"]
                    Sox = docE["Sox"]
            else:
                msgbox(msg="The Sox File Doesn't Exist")
                break

            while AccessOptions != "5":
                AccessOptions = ((buttonbox(msg=AccessOptionsMsg, title=ProgramTitle, choices=AccessOptionsChoices)).split("] ")[0]).replace("[","")

                if AccessOptions == "1":

                    userArray = []
                    headdeptArray = []
                    formdateArray = []
                    DeptArray = []
                    userArrayToSend = []
                    headdeptArrayToSend = []
                    formdateArrayToSend = []
                    DeptArrayToSend = []

                    SendAnswer = "No"
                    EndDate = " "
                    NoData = 1
                    ModifyExcel = 0

                    for x in range(2,(Sox.max_row)+1):
                        if Data.cell(row=x,column=1).value == None or len(Data.cell(row=x,column=1).value) == 1:
                            if Sox.cell(row=x,column=9).value == "Deleted":
                                NoData += 1
                                continue
                            else:
                                AccessFormFileExist = "No"
                        elif Data.cell(row=x,column=1).value != None and len(Data.cell(row=x,column=1).value) > 1:
                            AccessFormFileExist = "Yes"

                        if AccessFormFileExist == "Yes":
                            FirstName = Data.cell(row = x, column = 1).value
                            LastName = Data.cell(row = x, column = 2).value
                            EmployeeNum = Data.cell(row = x, column = 3).value
                            if EmployeeNum == "YYYY" or EmployeeNum == "XXXX":
                                EmployeeNum = "-"
                            TimeAttendanceNum = Data.cell(row = x, column = 4).value
                            if TimeAttendanceNum == "YYYY" or TimeAttendanceNum == "XXXX":
                                TimeAttendanceNum = "-"
                            SequenceNum = Data.cell(row = x, column = 5).value
                            Position = Data.cell(row = x, column = 6).value
                            Dept = Data.cell(row = x, column = 7).value
                            HeadDept1 = Data.cell(row = x, column = 8).value
                            HeadDept2 = Data.cell(row = x, column = 9).value
                            FormDate = Data.cell(row = x, column = 10).value
                            EndDate = Data.cell(row = x, column = 11).value
                            ItMan = Data.cell(row = x, column = 12).value
                            NetworkA = Data.cell(row = x, column = 13).value
                            NetworkId = Data.cell(row = x, column = 14).value
                            Program01A = Data.cell(row = x, column = 15).value
                            Program01Id = Data.cell(row = x, column = 16).value
                            OnQA = Data.cell(row = x, column = 17).value
                            OnQUserId = Data.cell(row = x, column = 18).value
                            OnQRights = Data.cell(row = x, column = 19).value
                            Program03A = Data.cell(row = x, column = 20).value
                            Program03User = Data.cell(row = x, column = 21).value
                            Program04A = Data.cell(row = x, column = 22).value
                            Program04User = Data.cell(row = x, column = 23).value
                            CompanyEmailA = Data.cell(row = x, column = 24).value
                            CompanyEmailAddress = Data.cell(row = x, column = 25).value
                            Program05A = Data.cell(row = x, column = 26).value
                            Program05User = Data.cell(row = x, column = 27).value
                            TsAccess = Data.cell(row = x, column = 28).value
                            TsUser = Data.cell(row = x, column = 29).value
                            PayRollAccess = Data.cell(row = x, column = 30).value
                            PayRollUser = Data.cell(row = x, column = 31).value
                            Time AttendanceA = Data.cell(row = x, column = 32).value
                            Time AttendanceUser = Data.cell(row = x, column = 33).value
                            Program07A = Data.cell(row = x, column = 34).value
                            Program07Rights = Data.cell(row = x, column = 35).value
                            Program07Card = Data.cell(row = x, column = 36).value
                            Program08A = Data.cell(row = x, column = 37).value
                            Program08AType = Data.cell(row = x, column = 38).value
                            Program08Id = Data.cell(row = x, column = 39).value
                            PersonalEmail = Data.cell(row = x, column = 41).value
                            Program10A = Data.cell(row = x, column = 42).value
                            Program10User = Data.cell(row = x, column = 43).value
                            Program11A = Data.cell(row = x, column = 44).value
                            Program11User = Data.cell(row = x, column = 45).value
                            SendAnswer = Data.cell(row = x, column = 48).value
                            Key01A = Data.cell(row = x, column = 49).value
                            Key01Type = Data.cell(row = x, column = 50).value
                            Program09A = Data.cell(row = x, column = 51).value
                            Program09Dept = Data.cell(row = x, column = 52).value
                            Program09User = Data.cell(row = x, column = 53).value
                            
                            if Sox.cell(row=x,column=9).value == "Leaving":
                                FirstName = Sox.cell(row=x,column=1).value
                                LastName = Sox.cell(row=x,column=2).value
                                EmployeeNum = Sox.cell(row=x,column=6).value
                                if EmployeeNum == "XXXX" or EmployeeNum == "YYYY":
                                    EmployeeNum = "-"
                                LeavingDate = Sox.cell(row=x,column=10).value
                                DateYForm = LeavingDate.split("/")[2]
                                DateMForm = LeavingDate.split("/")[1]
                                CorrectFormDate = LeavingDate.replace("/","-")
                                os.makedirs(FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)], exist_ok=True)
                                os.makedirs(FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated", exist_ok=True)

                                if charCounter(string = FirstName, char = " ") == 1:
                                    if charCounter(string = LastName, char = " ") == 1:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                    elif charCounter(string = LastName, char = " ") == 2:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "-" + CorrectFormDate + ".docx"
                                    elif charCounter(string = LastName, char = " ") == 3:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "_" + LastName.split(" ")[3] + "-" + CorrectFormDate + ".docx"
                                    else:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName + "-" + CorrectFormDate + ".docx"
                                elif charCounter(string = FirstName, char = " ") == 2:
                                    if charCounter(string = LastName, char = " ") == 1:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                    elif charCounter(string = LastName, char = " ") == 2:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "-" + CorrectFormDate + ".docx"
                                    elif charCounter(string = LastName, char = " ") == 3:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "_" + LastName.split(" ")[3] + "-" + CorrectFormDate + ".docx"
                                    else:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName + "-" + CorrectFormDate + ".docx"
                                elif charCounter(string = FirstName, char = " ") == 3:
                                    if charCounter(string = LastName, char = " ") == 1:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                    elif charCounter(string = LastName, char = " ") == 2:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "-" + CorrectFormDate + ".docx"
                                    elif charCounter(string = LastName, char = " ") == 3:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "_" + LastName.split(" ")[3] + "-" + CorrectFormDate + ".docx"
                                    else:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName + "-" + CorrectFormDate + ".docx"

                                else:
                                    if charCounter(string = LastName, char = " ") == 1:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                    else:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName + "_" + LastName + "-" + CorrectFormDate + ".docx"
                                
                                if os.path.exists(NameTerminatedFile) == True and os.path.isfile(NameTerminatedFile) == True:
                                    continue
                                else:
                                    terminatingUsers(FirstName, LastName, EmployeeNum, LeavingDate, NameTerminatedFile)
                                    continue
                        else:
                            if Sox.cell(row=x,column=9).value == "Deleted":
                                NoData += 1
                                continue
                            if Sox.cell(row=x,column=9).value == "Leaving":
                                FirstName = Sox.cell(row=x,column=1).value
                                LastName = Sox.cell(row=x,column=2).value
                                EmployeeNum = Sox.cell(row=x,column=6).value
                                if EmployeeNum == "XXXX" or EmployeeNum == "YYYY":
                                    EmployeeNum = "-"
                                LeavingDate = Sox.cell(row=x,column=10).value
                                DateYForm = LeavingDate.split("/")[2]
                                DateMForm = LeavingDate.split("/")[1]
                                CorrectFormDate = LeavingDate.replace("/","-")
                                os.makedirs(FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)], exist_ok=True)
                                os.makedirs(FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated", exist_ok=True)

                                if charCounter(string = FirstName, char = " ") == 1:
                                    if charCounter(string = LastName, char = " ") == 1:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                    elif charCounter(string = LastName, char = " ") == 2:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "-" + CorrectFormDate + ".docx"
                                    elif charCounter(string = LastName, char = " ") == 3:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "_" + LastName.split(" ")[3] + "-" + CorrectFormDate + ".docx"
                                    else:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName + "-" + CorrectFormDate + ".docx"
                                elif charCounter(string = FirstName, char = " ") == 2:
                                    if charCounter(string = LastName, char = " ") == 1:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                    elif charCounter(string = LastName, char = " ") == 2:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "-" + CorrectFormDate + ".docx"
                                    elif charCounter(string = LastName, char = " ") == 3:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "_" + LastName.split(" ")[3] + "-" + CorrectFormDate + ".docx"
                                    else:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName + "-" + CorrectFormDate + ".docx"
                                elif charCounter(string = FirstName, char = " ") == 3:
                                    if charCounter(string = LastName, char = " ") == 1:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                    elif charCounter(string = LastName, char = " ") == 2:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "-" + CorrectFormDate + ".docx"
                                    elif charCounter(string = LastName, char = " ") == 3:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "_" + LastName.split(" ")[3] + "-" + CorrectFormDate + ".docx"
                                    else:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName + "-" + CorrectFormDate + ".docx"

                                else:
                                    if charCounter(string = LastName, char = " ") == 1:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                    else:
                                        NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName + "_" + LastName + "-" + CorrectFormDate + ".docx"
                                
                                if os.path.exists(NameTerminatedFile) == True and os.path.isfile(NameTerminatedFile) == True:
                                    continue
                                else:
                                    terminatingUsers(FirstName, LastName, EmployeeNum, LeavingDate, NameTerminatedFile)
                                    continue
                            if Sox.cell(row=x,column=1).value == None:
                                NoData += 1
                                continue
                            
                            FirstName = Sox.cell(row = x, column = 1).value
                            LastName = Sox.cell(row = x, column = 2).value
                            EmployeeNum = Sox.cell(row = x, column = 6).value
                            if EmployeeNum == "XXXX" or EmployeeNum == "YYYY":
                                EmployeeNum = "-"
                            TimeAttendanceNum = Sox.cell(row = x, column = 5).value
                            if TimeAttendanceNum == "XXXX" or TimeAttendanceNum == "YYYY":
                                TimeAttendanceNum = "-"
                            PersonalEmail =  Sox.cell(row=x,column=3).value
                            SequenceNum = months[int((Sox.cell(row=x,column=10).value).split("/")[1])] + " - " + str(x-NoData)
                            Position = Sox.cell(row = x, column = 7).value
                            Dept = Sox.cell(row = x, column = 8).value
                            FormDate = Sox.cell(row = x, column = 10).value
                            
                            FirstName = replaceSpecialChars(FirstName)
                            LastName = replaceSpecialChars(LastName)

                            if LastName[:3] == "Van":
                                LastNameNether = True
                            if LastName[:2] == "De":
                                LastNameNether = True
                            else:
                                LastNameNether = False

                            if LastNameNether == False:
                                if charCounter(string = FirstName) >= 1:
                                    if charCounter(string = LastName) >= 1:
                                        if (len(FirstName.split(" ")[0] + LastName.split(" ")[0])) > 16:
                                            if (len(FirstName.split(" ")[0] + LastName.split(" ")[1])) > 16:
                                                if(len(FirstName.split(" ")[1] + LastName.split(" ")[0])) > 16:
                                                    if len(FirstName.split(" ")[1] + LastName.split(" ")[1]) > 16:
                                                        if len(FirstName.split(" ")[0][:1] + LastName.split(" ")[0]) > 16:
                                                            if len(FirstName.split(" ")[0][:1] + LastName.split(" ")[1]) > 16:
                                                                if len(FirstName.split(" ")[1][:1] + LastName.split(" ")[0]) > 16:
                                                                    if len(FirstName.split(" ")[1][:1] + LastName.split(" ")[1]) > 16:
                                                                        UserNameTM = enterbox(msg="Write the correct Username", default=UserNameTM)
                                                                    else:
                                                                        UserNameTM = ((FirstName.split(" ")[1])[:1] + LastName.split(" ")[1])[:16]
                                                                else:
                                                                    UserNameTM = ((FirstName.split(" ")[1])[:1] + LastName.split(" ")[0])[:16]
                                                            else:
                                                                UserNameTM = ((FirstName.split(" ")[0])[:1] + LastName.split(" ")[1])[:16]
                                                        else:
                                                            UserNameTM = ((FirstName.split(" ")[0])[:1] + LastName.split(" ")[0])[:16]
                                                    else:
                                                        UserNameTM = (FirstName.split(" ")[1] + LastName.split(" ")[1])[:16]
                                                else:
                                                    UserNameTM = (FirstName.split(" ")[1] + LastName.split(" ")[0])[:16]
                                            else:    
                                                UserNameTM = (FirstName.split(" ")[0] + LastName.split(" ")[1])[:16]
                                        else:
                                            UserNameTM = (FirstName.split(" ")[0] + LastName.split(" ")[0])[:16]
                                    else:
                                        if (len(FirstName.split(" ")[0] + LastName)) > 16:
                                            if (len(FirstName.split(" ")[1] + LastName)) > 16:
                                                UserNameChoices = [((FirstName.split(" ")[0])[:1] + LastName)[:16],((FirstName.split(" ")[1])[:1] + LastName)[:16]]
                                                UserNameTM = buttonbox(title=ProgramTitle, msg="\n\n\n \t\tChoose one Username", choices=UserNameChoices)
                                            else:
                                                UserNameTM = (FirstName.split(" ")[1] + LastName)[:16]
                                        else:
                                            UserNameTM = (FirstName.split(" ")[0] + LastName)[:16]
                                elif charCounter(string = LastName) >= 1:
                                    if (len(FirstName + LastName.split(" ")[0])) > 16:
                                        if (len(FirstName + LastName.split(" ")[1])) > 16:
                                            if (len(FirstName[:1] + LastName.split(" ")[0])) > 16:
                                                if (len(FirstName[:1] + LastName.split(" ")[1])) > 16:
                                                    UserNameTM = enterbox(msg="Write the correct Username", default=UserNameTM)
                                                else:
                                                    UserNameTM = (FirstName[:1] + LastName.split(" ")[1])[:16]
                                            else:
                                                UserNameTM = (FirstName[:1] + LastName.split(" ")[0])[:16]
                                        else:
                                            UserNameTM = (FirstName+ LastName.split(" ")[1])[:16]
                                    else:
                                        UserNameTM = (FirstName + LastName.split(" ")[0])[:16]
                                else:
                                    if (len(FirstName + LastName)) > 16:
                                        UserNameTM = (FirstName[:1] + LastName)[:16]
                                    else:
                                        UserNameTM = (FirstName + LastName)[:16]
                            else:
                                LastName = LastName.replace(" ", "")

                                UserNameTM = FirstName + LastName

                            VerificationUserName = boolbox(msg="Is " + UserNameTM + " an Accepted Username for " + FirstName + " " + LastName + "?")
                            if VerificationUserName == False:
                                CorrectionUserName = enterbox(msg="Write the correct Username", default=UserNameTM)
                                UserNameTM = CorrectionUserName[:16]

                            
                            if Dept == "Dept01":
                                HeadDept1 = "Example01"
                            elif Dept == "Dept02":
                                HeadDept1 = "Example02"
                            elif Dept == "Dept03":
                                HeadDept1 = "Example03"
                            elif Dept == "Dept04":
                                HeadDept1 = "Example04"
                            elif Dept == "Dept05":
                                HeadDept1 = "Example04"
                            else:
                                HeadDept1 =enterbox(msg="Enter the HoD Name", default="Head of Department")


                            HeadDept2 = " "

                            AccessFormFilemsg = "User: " + FirstName + " " + LastName + "\nUsername: " + UserNameTM + "\nDepartment: " + Dept + "\nPosition: " + Position + "\nHead of Department: " + HeadDept1
                            AccessFormFileFields = ["Network Access", "Intranet Access", "Program01 Access", "Program02 Access", "Program03 Access", "Program04 Access", "Company Email Access", "Program05 Access", "Terminal Server Access", "PayRoll Access", "Time Attendance Access", "Program07 Access", "Program08 Access", "Program09 Access", "Program10 Access", "Office Key Access"]

                            NetworkA = "No"
                            NetworkId = " "
                            Program01A = "No"
                            Program01Id = " "
                            Program02A = "No"
                            Program02UserId = " "
                            Program02Rights = " "
                            Program03A = "No"
                            Program03User = " "
                            Program04A = "No"
                            Program04User = " "
                            CompanyEmailA = "No"
                            CompanyEmailAddress = " "
                            Program05A = "No"
                            Program05User = " "
                            Program06Access = "No"
                            Program06User = " "
                            PayRollAccess = "No"
                            PayRollUser = " "
                            TimeAttendanceA = "No"
                            TimeAttendanceUser = " "
                            Program07A = "No"
                            Program07Rights = " "
                            Program07Card = " "
                            Program08A = "No"
                            Program08AType = " "
                            Program08Id = " "
                            Program09A = "No"
                            Program09User = " "
                            Program10A = "No"
                            Program10User = " "
                            SendAnswer = "No"
                            Key01A = "No"
                            Key01Type = " "
                            Program11A = "No"
                            Program11Dept = " "
                            Program11User = " "

                            ItMan = "IT Manager"


                            AccessFormChoices = multchoicebox(title=ProgramTitle, msg=AccessFormFilemsg, choices=AccessFormFileFields)

                            for z in AccessFormChoices:
                                if z == "Network Access":
                                    NetworkA = "Yes"
                                    NetworkId = "Domain\\" + UserNameTM
                                elif z == "Intranet Access":
                                    Program01A = "Yes"
                                    Program01Id = UserNameTM
                                elif z == "Program01 Access":
                                    OnQA = "Yes"
                                    OnQUserId = UserNameTM.upper()
                                    if Dept == "Example01":
                                        OnQRights = "Example01 Staff"
                                    elif Dept == "Example02":
                                        OnQRights = "Example02 Staff"
                                    elif Dept == "Example03":
                                        OnQRights = "Example03 Staff"
                                elif z == "Program02 Access":
                                    Program11A = "Yes"
                                    Program11User = UserNameTM.upper()
                                elif z == "Program03 Access":
                                    Program03 = "Yes"
                                    Program03User = "Domain\\" + UserNameTM
                                elif z == "Program04 Access":
                                    Program04A = "Yes"
                                    Program04User = (UserNameTM[:1] + UserNameTM[len(FirstName):]).upper()
                                elif z == "Company Email Access":
                                    CompanyEmailA = "Yes"
                                    if charCounter(string = FirstName) == 1:
                                        if charCounter(string = LastName) == 1:
                                            CompanyEmailAddress = FirstName.split(" ")[0] + "." + LastName.split(" ")[0] + "@Company.com"
                                        else:
                                            CompanyEmailAddress = FirstName.split(" ")[0] + "." + LastName + "@Company.com"
                                    else:
                                        if charCounter(string = LastName) == 1:
                                            CompanyEmailAddress = FirstName + "." + LastName.split(" ")[0] + "@Company.com"
                                        else:
                                            CompanyEmailAddress = FirstName + "." + LastName + "@Company.com"
                                elif z == "Program05 Access":
                                    Program05A = "Yes"
                                    Program05User = UserNameTM
                                elif z == "Program06 Access":
                                    Program06Access = "Yes"
                                    Program06User = UserNameTM
                                elif z == "PayRoll Access":
                                    PayRollAccess = "Yes"
                                    PayRollUser = UserNameTM
                                elif z == "Time Attendance Access":
                                    TimeAttendanceA = "Yes"
                                    TimeAttendanceUser = UserNameTM
                                elif z == "Program07 Access":
                                    Program07A = "Yes"
                                    Program07Card = "XXXXXXXXX"
                                    Program07 = Position
                                elif z == "Program08 Access":
                                    Program08A  ="Yes"
                                    Program08Id = UserNameTM
                                    Program08AType = Dept
                                elif z == "Program09 Access":
                                    Program09A = "Yes"
                                    Program09User = UserNameTM
                                    Program09Dept = Dept
                                elif z == "Program10 Access":
                                    Program10A = "Yes"
                                    Program10User = "CompanyXXX"
                                elif z == "Key01 Access":
                                    Key01A = "Yes"
                                    if Dept == "Dept01":
                                        Key01Type = "Example01"
                                    if Dept == "Dept02":
                                        Key01Type = "Example02"
                                    if Dept == "Dept03":
                                        Key01Type = "Example03"
                                    if Dept == "Dept04":
                                        Key01Type = "Example04"
                                    if Dept == "Dept05":
                                        Key01Type = "Example05"
                                    else:
                                        Key01Type = enterbox(msg="Write the Key Name", default="DESPXXX")

                            Data.cell(row=x,column=1, value=FirstName)
                            Data.cell(row=x,column=2, value=LastName)
                            Data.cell(row=x,column=3, value=EmployeeNum)
                            Data.cell(row=x,column=4, value=TimeAttendanceNum)
                            Data.cell(row=x,column=5, value=SequenceNum)
                            Data.cell(row=x,column=6, value=Position)
                            Data.cell(row=x,column=7, value=Dept)
                            Data.cell(row=x,column=8, value=HeadDept1)
                            Data.cell(row=x,column=9, value=HeadDept2)
                            Data.cell(row=x,column=10, value=FormDate)
                            Data.cell(row=x,column=11, value=EndDate)
                            Data.cell(row=x,column=12, value=ItMan)
                            Data.cell(row=x,column=13, value=NetworkA)
                            Data.cell(row=x,column=14, value=NetworkId)
                            Data.cell(row=x,column=15, value=Program01A)
                            Data.cell(row=x,column=16, value=Program01Id)
                            Data.cell(row=x,column=17, value=Program02A)
                            Data.cell(row=x,column=18, value=Program02UserId)
                            Data.cell(row=x,column=19, value=Program02Rights)
                            Data.cell(row=x,column=20, value=Program03A)
                            Data.cell(row=x,column=21, value=Program03User)
                            Data.cell(row=x,column=22, value=Program04A)
                            Data.cell(row=x,column=23, value=Program04User)
                            Data.cell(row=x,column=24, value=CompanyEmailA)
                            Data.cell(row=x,column=25, value=CompanyEmailAddress)
                            Data.cell(row=x,column=26, value=Program05A)
                            Data.cell(row=x,column=27, value=Program05User)
                            Data.cell(row=x,column=28, value=Program06Access)
                            Data.cell(row=x,column=29, value=Program06User)
                            Data.cell(row=x,column=30, value=PayRollAccess)
                            Data.cell(row=x,column=31, value=PayRollUser)
                            Data.cell(row=x,column=32, value=TimeAttendanceA)
                            Data.cell(row=x,column=33, value=TimeAttendanceUser)
                            Data.cell(row=x,column=34, value=Program07A)
                            Data.cell(row=x,column=35, value=Program07Rights)
                            Data.cell(row=x,column=36, value=Program07Card)
                            Data.cell(row=x,column=37, value=Program08A)
                            Data.cell(row=x,column=38, value=Program08AType)
                            Data.cell(row=x,column=39, value=Program08Id)
                            Data.cell(row=x,column=41, value=PersonalEmail)
                            Data.cell(row=x,column=42, value=Program10A)
                            Data.cell(row=x,column=43, value=Program10User)
                            Data.cell(row=x,column=44, value=Program11A)
                            Data.cell(row=x,column=45, value=Program11User)
                            Data.cell(row=x,column=48, value=SendAnswer)
                            Data.cell(row=x,column=49, value=Key01A)
                            Data.cell(row=x,column=50, value=Key01Type)
                            Data.cell(row=x,column=51, value=Program09A)
                            Data.cell(row=x,column=52, value=Program09Dept)
                            Data.cell(row=x,column=53, value=Program09User)

                            if Sox.cell(row=x,column=9).value == "Starting":
                                Deleted = 0
                                Leaving = 0

                            DateYForm = str(FormDate.split("/")[2])
                    else:
                        if ModifyExcel == 1:
                            docE.save(TMPAccessExcel2)
                        elif ModifyExcel == 0:
                            docE.save(TMPAccessExcel)

                    DeletedLine = 0
                    for y in range(2,(Data.max_row + DeletedLine)+1):

                        if ModifyExcel == 1:
                             docW = Document(AccessFormFile)
                             docE = load_workbook(TMPAccessExcel2)
                             Data = docE["Users"]
                             Sox = docE["Sox"]
                        elif ModifyExcel == 0:
                            docW = Document(AccessFormFile)
                            docE = load_workbook(TMPAccessExcel)
                            Data = docE["Users"]
                            Sox = docE["Sox"]

                        if Data.cell(row = y, column = 1).value == None:
                            DeletedLine = DeletedLine + 1
                            continue

                        FirstName = Data.cell(row = y, column = 1).value
                        LastName = Data.cell(row = y, column = 2).value
                        EmployeeNum = Data.cell(row = y, column = 3).value
                        if EmployeeNum == "XXXX" or EmployeeNum == "YYYY":
                            EmployeeNum = "-"
                        TimeAttendanceNum = Data.cell(row = y, column = 4).value
                        if TimeAttendanceNum == "XXXX" or TimeAttendanceNum == "YYYY":
                            TimeAttendanceNum = "-"
                        SequenceNum = Data.cell(row = y, column = 5).value
                        Position = Data.cell(row = y, column = 6).value
                        Dept = Data.cell(row = y, column = 7).value
                        HeadDept1 = Data.cell(row = y, column = 8).value
                        HeadDept2 = Data.cell(row = y, column = 9).value
                        FormDate = Data.cell(row = y, column = 10).value
                        EndDate = Data.cell(row = y, column = 11).value
                        ItMan = Data.cell(row = y, column = 12).value
                        NetworkA = Data.cell(row = y, column = 13).value
                        NetworkId = Data.cell(row = y, column = 14).value
                        Program01A = Data.cell(row = y, column = 15).value
                        Program01Id = Data.cell(row = y, column = 16).value
                        Program02A = Data.cell(row = y, column = 17).value
                        Program02UserId = Data.cell(row = y, column = 18).value
                        Program02Rights = Data.cell(row = y, column = 19).value
                        Program03A = Data.cell(row = y, column = 20).value
                        Program03User = Data.cell(row = y, column = 21).value
                        Program04A = Data.cell(row = y, column = 22).value
                        Program04User = Data.cell(row = y, column = 23).value
                        CompanyEmailA = Data.cell(row = y, column = 24).value
                        CompanyEmailAddress = Data.cell(row = y, column = 25).value
                        Program05A = Data.cell(row = y, column = 26).value
                        Program05User = Data.cell(row = y, column = 27).value
                        Program06Access = Data.cell(row = y, column = 28).value
                        Program06User = Data.cell(row = y, column = 29).value
                        PayRollAccess = Data.cell(row = y, column = 30).value
                        PayRollUser = Data.cell(row = y, column = 31).value
                        TimeAttendanceA = Data.cell(row = y, column = 32).value
                        TimeAttendanceUser = Data.cell(row = y, column = 33).value
                        Program07A = Data.cell(row = y, column = 34).value
                        Program07Rights = Data.cell(row = y, column = 35).value
                        Program07Card = Data.cell(row = y, column = 36).value
                        Program08A = Data.cell(row = y, column = 37).value
                        Program08AType = Data.cell(row = y, column = 38).value
                        Program08Id = Data.cell(row = y, column = 39).value
                        PersonalEmail = Data.cell(row = y, column = 41).value
                        Program10A = Data.cell(row = y, column = 42).value
                        Program10User = Data.cell(row = y, column = 43).value
                        Program11A = Data.cell(row = y, column = 44).value
                        Program11User = Data.cell(row = y, column = 45).value
                        SendAnswer = Data.cell(row = y, column = 48).value
                        Key01A = Data.cell(row = y, column = 49).value
                        Key01Type = Data.cell(row = y, column = 50).value
                        Program09A = Data.cell(row = y, column = 51).value
                        Program09Dept = Data.cell(row = y, column = 52).value
                        Program09User = Data.cell(row = y, column = 53).value

                        docW._body.clear_content()
                        #User information Table
                        Templatetable1 = docW.add_table(rows=4, cols=5, style="Table Grid")
                        #Space between Templatetable1 and Templatetable2
                        docW.add_paragraph().add_run(" ").font.size = Pt(8)
                        #Program access Table
                        Templatetable2 = docW.add_table(rows=16, cols=5, style="Table Grid")
                        #Space between Templatetable2 and Fottertable1
                        docW.add_paragraph(" ", style = "Intense Quote")
                        #Approval Signature zone
                        Fottertable1 = docW.add_table(rows=2, cols=6)
                        #Page break to have space for the Templatetable2
                        # docW.add_paragraph().add_run(" ").font.size = Pt(20)
                        # docW.add_paragraph().add_run(" ").font.size = Pt(20)
                        #Table for the information
                        Fottertable2 = docW.add_table(rows=4, cols=1)
                        #Space between text info (Fottertable2) and Fottertable3
                        docW.add_paragraph(" ", style = "Intense Quote")
                        #zone for Date and Employee signature
                        Fottertable3 = docW.add_table(rows=1, cols=6)

                        """
                        FIRST TABLE/FIRST TABLE/FIRST TABLE
                        """

                        Templatetable1.rows[0].cells[0].text = "Team Member Name:"
                        Templatetable1.rows[0].cells[0].paragraphs[0].runs[0].bold = True
                        Templatetable1.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable1.rows[0].cells[1].text = FirstName + " " + LastName
                        Templatetable1.rows[0].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable1.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable1.rows[0].cells[2].text = "Employee Num."
                        Templatetable1.rows[0].cells[2].paragraphs[0].runs[0].bold = True
                        Templatetable1.rows[0].cells[2].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable1.rows[0].cells[3].text = EmployeeNum
                        Templatetable1.rows[0].cells[3].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable1.rows[0].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable1.rows[0].cells[4].text = "Time Attendance Num"
                        Templatetable1.rows[0].cells[4].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable1.rows[0].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable1.rows[1].cells[0].text = "Sequence Number:"
                        Templatetable1.rows[1].cells[0].paragraphs[0].runs[0].bold = True
                        Templatetable1.rows[1].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable1.rows[1].cells[1].text = SequenceNum
                        Templatetable1.rows[1].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable1.rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable1.rows[1].cells[2].text = "Position / Dept.:"
                        Templatetable1.rows[1].cells[2].paragraphs[0].runs[0].bold = True
                        Templatetable1.rows[1].cells[2].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable1.rows[1].cells[3].text = Position + " - " + Dept
                        Templatetable1.rows[1].cells[3].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable1.rows[1].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        a = Templatetable1.rows[1].cells[3]
                        b = Templatetable1.rows[1].cells[4]
                        a.merge(b)
                        Templatetable1.rows[2].cells[0].text = "Effective Date of Form:"
                        Templatetable1.rows[2].cells[0].paragraphs[0].runs[0].bold = True
                        Templatetable1.rows[2].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable1.rows[2].cells[1].text = FormDate
                        Templatetable1.rows[2].cells[1].paragraphs[0].runs[0].font.size = Pt(8)
                        Templatetable1.rows[2].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable1.rows[2].cells[2].text = "Terminate Date:"
                        Templatetable1.rows[2].cells[2].paragraphs[0].runs[0].bold = True
                        Templatetable1.rows[2].cells[2].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable1.rows[2].cells[3].text = EndDate
                        Templatetable1.rows[2].cells[3].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable1.rows[2].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        a = Templatetable1.rows[2].cells[3]
                        b = Templatetable1.rows[2].cells[4]
                        a.merge(b)
                        Templatetable1.rows[3].cells[0].text = "Personal Email:"
                        Templatetable1.rows[3].cells[0].paragraphs[0].runs[0].bold = True
                        Templatetable1.rows[3].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                        # a = Templatetable1.rows[3].cells[1]
                        # b = Templatetable1.rows[3].cells[2]
                        # a.merge(b)
                        Templatetable1.rows[3].cells[1].text = PersonalEmail.lower()
                        Templatetable1.rows[3].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable1.rows[3].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable1.rows[3].cells[3].text = "Office Key:"
                        Templatetable1.rows[3].cells[3].paragraphs[0].runs[0].bold = True
                        Templatetable1.rows[3].cells[3].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable1.rows[3].cells[4].text = Key01Type
                        if len(Key01Type) > 25:
                            Templatetable1.rows[3].cells[4].paragraphs[0].runs[0].font.size = Pt(7)
                        else:
                            Templatetable1.rows[3].cells[4].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable1.rows[3].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        a = Templatetable1.rows[3].cells[1]
                        b = Templatetable1.rows[3].cells[2]
                        a.merge(b)

                        """
                        SECOND TABLE/SECOND TABLE/SECOND TABLE
                        """
                        Templatetable2.rows[0].cells[0].text = "Network access:"
                        Templatetable2.rows[0].cells[1].text = NetworkA
                        Templatetable2.rows[0].cells[2].text = "Username:"
                        Templatetable2.rows[0].cells[3].text = NetworkId
                        Templatetable2.rows[1].cells[0].text = "Intranet access:"
                        Templatetable2.rows[1].cells[1].text = Program01A
                        Templatetable2.rows[1].cells[2].text = "Username:"
                        Templatetable2.rows[1].cells[3].text = Program01Id
                        Templatetable2.rows[2].cells[0].text = "Program01 access:"
                        Templatetable2.rows[2].cells[1].text = Program02A
                        Templatetable2.rows[2].cells[2].text = "Username:"
                        Templatetable2.rows[2].cells[3].text = Program02UserId
                        Templatetable2.rows[3].cells[0].text = "Program01 rights:"
                        Templatetable2.rows[3].cells[1].text = Program02Rights
                        Templatetable2.rows[4].cells[0].text = "Program02 access:"
                        Templatetable2.rows[4].cells[1].text = Program11A
                        Templatetable2.rows[4].cells[2].text = "Username:"
                        Templatetable2.rows[4].cells[3].text = Program11User
                        Templatetable2.rows[5].cells[0].text = "Program03 access:"
                        Templatetable2.rows[5].Cells[1].text = Program03A
                        Templatetable2.rows[5].cells[2].text = "Username:"
                        Templatetable2.rows[5].cells[3].text = Program03User
                        Templatetable2.rows[6].cells[0].text = "Program04 access:"
                        Templatetable2.rows[6].cells[1].text = Program04A
                        Templatetable2.rows[6].cells[2].text = "Username:"
                        Templatetable2.rows[6].cells[3].text = Program04User
                        Templatetable2.rows[7].cells[0].text = "Company Email Access:"
                        Templatetable2.rows[7].cells[1].text = CompanyEmailA
                        Templatetable2.rows[7].cells[2].text = "Username:"
                        Templatetable2.rows[7].cells[3].text = CompanyEmailAddress
                        Templatetable2.rows[8].cells[0].text = "Program05 access:"
                        Templatetable2.rows[8].cells[1].text = Program05 Access
                        Templatetable2.rows[8].cells[2].text = "Username:"
                        Templatetable2.rows[8].cells[3].text = Program05 User
                        Templatetable2.rows[9].cells[0].text = "Payroll Apps. access:"
                        Templatetable2.rows[9].cells[1].text = PayRollAccess
                        Templatetable2.rows[9].cells[2].text = "Username:"
                        Templatetable2.rows[9].cells[3].text = PayRollUser
                        Templatetable2.rows[10].cells[0].text = "Time Attendance access:"
                        Templatetable2.rows[10].cells[1].text = TimeAttendanceA
                        Templatetable2.rows[10].cells[2].text = "Username:"
                        Templatetable2.rows[10].cells[3].text = TimeAttendanceUser
                        Templatetable2.rows[11].cells[0].text = "Program07 access:"
                        Templatetable2.rows[11].cells[1].text = Program07A
                        Templatetable2.rows[11].cells[2].text = Program07Rights
                        Templatetable2.rows[11].cells[3].text = "Program07 Card:"
                        Templatetable2.rows[11].cells[4].text = Program07Card
                        Templatetable2.rows[12].cells[0].text = "Program05 access:"
                        Templatetable2.rows[12].cells[1].text = Program05A
                        Templatetable2.rows[12].cells[2].text = "Username:"
                        Templatetable2.rows[12].cells[3].text = Program05User
                        Templatetable2.rows[13].cells[0].text = "Program08 access:"
                        Templatetable2.rows[13].cells[1].text = Program08A
                        Templatetable2.rows[13].cells[2].text = "Username:"
                        Templatetable2.rows[13].cells[3].text = Program08Id
                        Templatetable2.rows[13].cells[4].text = Program08AType
                        Templatetable2.rows[14].cells[0].text = "Program09 Access:"
                        Templatetable2.rows[14].cells[1].text = Program09A
                        Templatetable2.rows[14].cells[2].text = "Username:"
                        Templatetable2.rows[14].cells[3].text = Program09User
                        Templatetable2.rows[14].cells[4].text = Program09Dept
                        Templatetable2.rows[15].cells[0].text = "Program10 Access:"
                        Templatetable2.rows[15].cells[1].text = Program10A
                        Templatetable2.rows[15].cells[2].text = "Username:"
                        Templatetable2.rows[15].cells[3].text = Program10User
                        Templatetable2.rows[0].cells[0].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[0].cells[2].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[1].cells[0].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[1].cells[2].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[2].cells[0].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[2].cells[2].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[3].cells[0].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[4].cells[0].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[4].cells[2].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[5].cells[0].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[5].cells[2].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[6].cells[0].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[6].cells[2].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[7].cells[0].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[7].cells[2].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[8].cells[0].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[8].cells[2].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[9].cells[0].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[9].cells[2].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[10].cells[0].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[10].cells[2].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[11].cells[0].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[11].cells[3].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[12].cells[0].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[12].cells[2].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[13].cells[0].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[13].cells[2].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[14].cells[0].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[14].cells[2].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[15].cells[0].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[15].cells[2].paragraphs[0].runs[0].bold = True
                        Templatetable2.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[0].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[0].cells[2].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[0].cells[3].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[1].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[1].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[1].cells[2].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[1].cells[3].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[2].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[2].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[2].cells[2].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[2].cells[3].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[3].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[3].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[4].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[4].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[4].cells[2].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[4].cells[3].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[5].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[5].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[5].cells[2].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[5].cells[3].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[6].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[6].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[6].cells[2].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[6].cells[3].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[7].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[7].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[7].cells[2].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[7].cells[3].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[8].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[8].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[8].cells[2].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[8].cells[3].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[9].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[9].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[9].cells[2].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[9].cells[3].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[10].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[10].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[10].cells[2].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[10].cells[3].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[11].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[11].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[11].cells[2].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[11].cells[3].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[11].cells[4].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[12].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[12].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[12].cells[2].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[12].cells[3].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[13].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[13].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[13].cells[2].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[13].cells[3].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[13].cells[4].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[14].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[14].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[14].cells[2].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[14].cells[3].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[14].cells[4].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[15].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[15].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[15].cells[2].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[15].cells[3].paragraphs[0].runs[0].font.size = Pt(9)
                        Templatetable2.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[0].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[1].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[2].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[2].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[3].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[4].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[4].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[5].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[5].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[6].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[6].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[7].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[7].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[8].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[8].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[9].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[9].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[10].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[10].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[11].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[11].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[12].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[12].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[13].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[13].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[13].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[14].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[14].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[14].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[15].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Templatetable2.rows[15].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        a = Templatetable2.rows[0].cells[3]
                        b = Templatetable2.rows[0].cells[4]
                        a.merge(b)
                        a = Templatetable2.rows[1].cells[3]
                        b = Templatetable2.rows[1].cells[4]
                        a.merge(b)
                        a = Templatetable2.rows[2].cells[3]
                        b = Templatetable2.rows[2].cells[4]
                        a.merge(b)
                        a = Templatetable2.rows[3].cells[1]
                        b = Templatetable2.rows[3].cells[2]
                        c = Templatetable2.rows[3].cells[3]
                        d = Templatetable2.rows[3].cells[4]
                        a.merge(b)
                        a.merge(c)
                        a.merge(d)
                        a = Templatetable2.rows[4].cells[3]
                        b = Templatetable2.rows[4].cells[4]
                        a.merge(b)
                        a = Templatetable2.rows[5].cells[3]
                        b = Templatetable2.rows[5].cells[4]
                        a.merge(b)
                        a = Templatetable2.rows[6].cells[3]
                        b = Templatetable2.rows[6].cells[4]
                        a.merge(b)
                        a = Templatetable2.rows[7].cells[3]
                        b = Templatetable2.rows[7].cells[4]
                        a.merge(b)
                        a = Templatetable2.rows[8].cells[3]
                        b = Templatetable2.rows[8].cells[4]
                        a.merge(b)
                        a = Templatetable2.rows[9].cells[3]
                        b = Templatetable2.rows[9].cells[4]
                        a.merge(b)
                        a = Templatetable2.rows[10].cells[3]
                        b = Templatetable2.rows[10].cells[4]
                        a.merge(b)
                        a = Templatetable2.rows[12].cells[3]
                        b = Templatetable2.rows[12].cells[4]
                        a.merge(b)

                        """
                        THIRD TABLE/THIRD TABLE/THIRD TABLE
                        """
                        Fottertable1.rows[0].cells[0].text = "Department Head: "
                        Fottertable1.rows[0].cells[1].text = HeadDept1
                        Fottertable1.rows[0].cells[0].paragraphs[0].runs[0].italic = True
                        Fottertable1.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(8)
                        Fottertable1.rows[0].cells[1].paragraphs[0].runs[0].bold = True
                        Fottertable1.rows[0].cells[1].paragraphs[0].runs[0].font.size = Pt(8)
                        a = Fottertable1.rows[0].cells[0]
                        b = Fottertable1.rows[0].cells[1]
                        a.merge(b)

                        Fottertable1.rows[0].cells[2].text = "Cross-Training Department Head: "
                        if len(HeadDept2) <= 1:
                            Fottertable1.rows[0].cells[3].text = "Not Available"
                        else:
                            Fottertable1.rows[0].cells[3].text = HeadDept2
                        Fottertable1.rows[0].cells[2].paragraphs[0].runs[0].italic = True
                        Fottertable1.rows[0].cells[2].paragraphs[0].runs[0].font.size = Pt(8)
                        Fottertable1.rows[0].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        Fottertable1.rows[0].cells[3].paragraphs[0].runs[0].bold = True
                        Fottertable1.rows[0].cells[3].paragraphs[0].runs[0].font.size = Pt(8)
                        Fottertable1.rows[0].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        a = Fottertable1.rows[0].cells[2]
                        b = Fottertable1.rows[0].cells[3]
                        a.merge(b)

                        Fottertable1.rows[0].cells[4].text = "Information Systems Manager: "
                        Fottertable1.rows[0].cells[5].text = ItMan
                        Fottertable1.rows[0].cells[4].paragraphs[0].runs[0].italic = True
                        Fottertable1.rows[0].cells[4].paragraphs[0].runs[0].font.size = Pt(8)
                        Fottertable1.rows[0].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        Fottertable1.rows[0].cells[5].paragraphs[0].runs[0].bold = True
                        Fottertable1.rows[0].cells[5].paragraphs[0].runs[0].font.size = Pt(8)
                        Fottertable1.rows[0].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        a = Fottertable1.rows[0].cells[4]
                        b = Fottertable1.rows[0].cells[5]
                        a.merge(b)

                        Fottertable1.rows[1].cells[0].text = FormDate
                        Fottertable1.rows[1].cells[0].paragraphs[0].runs[0].italic = True
                        Fottertable1.rows[1].cells[0].paragraphs[0].runs[0].font.size = Pt(8)
                        if len(HeadDept2) <= 1:
                            Fottertable1.rows[1].cells[2].text = ""
                        else:
                            Fottertable1.rows[1].cells[2].text = FormDate
                            Fottertable1.rows[1].cells[2].paragraphs[0].runs[0].italic = True
                            Fottertable1.rows[1].cells[2].paragraphs[0].runs[0].font.size = Pt(8)
                        Fottertable1.rows[1].cells[5].paragraphs[0].add_run().add_picture("Files\\SignatureExample.png", width = Inches(1.40), height = Inches(0.96))
                        Fottertable1.rows[1].cells[4].text = FormDate
                        Fottertable1.rows[1].cells[4].paragraphs[0].runs[0].italic = True
                        Fottertable1.rows[1].cells[4].paragraphs[0].runs[0].font.size = Pt(8)
                        Fottertable1.rows[1].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        Fottertable1.rows[1].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        """
                        FOURTH TABLE/FOURTH TABLE/FOURTH TABLE
                        """
                        Fottertable2.rows[0].cells[0].text = "\nThe Username allows the undersigned to access to system functions designed for his/her job profile. Passwords and security codes must not be disclosed to any employees or other individuals since it would enable them to have access to sensitive hotel data. By signing this form, I acknowledge the Company requirement not to disclose data or procedures for accessing data to which I have no access or knowledge. Any breach of security or improper use of the systems will be ground for dismissal.\nI have read and understood the above and acknowledge the responsibility for my Username, password or any other security code. I will use them only for proper discharge of my responsibilities and for no other purpose.\nIf I forget my Username, password or security code, I will contact the Information Systems Manager (ISM) immediately. I will also contact the ISM if anyone learns of my password or security code."
                        Fottertable2.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(7)
                        Fottertable2.rows[0].cells[0].paragraphs[0].runs[0].italic = True
                        Fottertable2.rows[1].cells[0].text = "Key system access authorization"
                        Fottertable2.rows[1].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                        Fottertable2.rows[1].cells[0].paragraphs[0].runs[0].bold = True
                        Fottertable2.rows[2].cells[0].text = "All Team Members who have been authorized to generate / make guest keys must do so while maintaining the integrity of the guest and the Company Corporation. Below is a list of guidelines to be followed:"
                        Fottertable2.rows[2].cells[0].paragraphs[0].runs[0].font.size = Pt(7)
                        Fottertable2.rows[2].cells[0].paragraphs[0].runs[0].italic = True
                        Fottertable2.rows[3].cells[0].text = "\t * Always verify the identity of the subject for who you are making a key. Examine the subjects ID and confirm room authorization to the room. When in doubt, call Security.\n\t * Never make / generate a key for a non-guest. When in doubt, call Security.\n\t * Never make more than the requested number of keys. Once a key/keys have been made, keep it/them secured.\n\t * Never divulge your access code to anyone. If you believe that your code has been compromised, contact the Information Systems Manager."
                        Fottertable2.rows[3].cells[0].paragraphs[0].runs[0].font.size = Pt(7)
                        Fottertable2.rows[3].cells[0].paragraphs[0].runs[0].italic = True

                        """
                        FIFTH TABLE/FIFTH TABLE/FIFTH TABLE
                        """
                        Fottertable3.rows[0].cells[0].text = "Date: "
                        Fottertable3.rows[0].cells[1].text = FormDate
                        Fottertable3.rows[0].cells[0].paragraphs[0].runs[0].italic = True
                        Fottertable3.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(8)
                        Fottertable3.rows[0].cells[1].paragraphs[0].runs[0].bold = True
                        Fottertable3.rows[0].cells[1].paragraphs[0].runs[0].font.size = Pt(8)
                        a = Fottertable3.rows[0].cells[1]
                        b = Fottertable3.rows[0].cells[2]
                        a.merge(b)

                        Fottertable3.rows[0].cells[3].text = "Employee: "
                        Fottertable3.rows[0].cells[4].text = FirstName + " " + LastName
                        Fottertable3.rows[0].cells[3].paragraphs[0].runs[0].italic = True
                        Fottertable3.rows[0].cells[3].paragraphs[0].runs[0].font.size = Pt(8)
                        Fottertable3.rows[0].cells[4].paragraphs[0].runs[0].bold = True
                        Fottertable3.rows[0].cells[4].paragraphs[0].runs[0].font.size = Pt(8)
                        Fottertable3.rows[0].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        a = Fottertable3.rows[0].cells[4]
                        b = Fottertable3.rows[0].cells[5]
                        a.merge(b)

                        DateYForm = str(FormDate.split("/")[2])
                        DateMForm = str(FormDate.split("/")[1])
                        CorrectFormDate = FormDate.replace("/","-")
                        os.makedirs(FormFiles + "\\" + DateYForm, exist_ok=True)
                        os.makedirs(FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)], exist_ok=True)

                        if AccessFormFile == "Files\\NetworkUserFormLand.docx":
                            NameFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] +  "\\Company_" + FirstName + "_" + LastName + "-" + CorrectFormDate + "_Landscape_Mode" + ".docx"
                        elif AccessFormFile == "Files\\NetworkUserForm.docx":
                            if charCounter(string = FirstName, char = " ") == 1:
                                if charCounter(string = LastName, char = " ") == 1:
                                    NameFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] +  "\\Company_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                else:
                                    NameFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] +  "\\Company_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName + "-" + CorrectFormDate + ".docx"
                            else:
                                if charCounter(string = LastName, char = " ") == 1:
                                    NameFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] +  "\\Company_" + FirstName + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                else:
                                    NameFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] +  "\\Company_" + FirstName + "_" + LastName + "-" + CorrectFormDate + ".docx"

                        if SendAnswer == "No":
                            docW.save(NameFile)
                            covx_to_pdf(NameFile)
                        elif SendAnswer == "Yes":
                            UserNameToEnter = FirstName + " " + LastName
                            userArray.append(UserNameToEnter)
                            headdeptArray.append(HeadDept1)
                            formdateArray.append(FormDate)
                            DeptArray.append(Dept)

                            docW.save(NameFile)
                            covx_to_pdf(NameFile)
                    else:
                        if len(userArray) < 0:
                            break
                        elif len(userArray) >= 1:
                            userArrayToSend = multchoicebox(msg="Select Users to send the form", title=ProgramTitle, choices = userArray)
                            if len(userArrayToSend) != None:
                                headdeptArrayToSend[:] = headdeptArray
                                formdateArrayToSend[:] = formdateArray
                                DeptArrayToSend[:] = DeptArray
                                reverseuserArrayToSend = list(reversed(userArrayToSend))

                                for x in reverseuserArrayToSend:
                                    UserIndex = userArray.index(x)
                                    headdeptArray.remove(headdeptArrayToSend[UserIndex])
                                    formdateArray.remove(formdateArrayToSend[UserIndex])
                                    DeptArray.remove(DeptArrayToSend[UserIndex])
                                for x in range(0,len(headdeptArray)):
                                    headdeptArrayToSend.remove(headdeptArray[x])
                                    formdateArrayToSend.remove(formdateArray[x])
                                    DeptArrayToSend.remove(DeptArray[x])
                                for x in range(0,len(userArrayToSend)):
                                    dataCheck = []
                                    CorrectFormDate = formdateArrayToSend[x].replace("/","-")
                                    DateMForm = formdateArrayToSend[x].split("/")[1]
                                    DateYForm = formdateArrayToSend[x].split("/")[2]
                                    SubjectMsg = "Sox User Access Form " + months[int(DateMForm)] + " " + DateYForm + " - " + userArrayToSend[x].capitalize()
                                    NameFilePDF = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] +  "\\Company_" + (userArrayToSend[x].capitalize()).replace(" ","_") + "-" + CorrectFormDate + ".pdf"
                                    NameFileDocx = "Company_" + (userArrayToSend[x].capitalize()).replace(" ","_") + "-" + CorrectFormDate + ".docx"
                                    NameFilePDF1 = NameFileDocx.replace(".docx",".pdf")

                                    if charCounter(string = headdeptArrayToSend[x], char = " ") == 1:
                                        HeadDeptEmail = headdeptArrayToSend[x].replace(" ",".") + "@Company.com"
                                    elif charCounter(string = headdeptArrayToSend[x], char = " ") == 2:
                                        HeadDeptEmail = headdeptArrayToSend[x].split(" ")[0] + headdeptArrayToSend[x].split(" ")[1] + "." + headdeptArrayToSend[x].split(" ")[2] + "@Company.com"
                                    dataCheck.append(userArrayToSend[x])
                                    dataCheck.append(formdateArrayToSend[x])
                                    dataCheck.append(DeptArrayToSend[x])
                                    dataCheck.append(HeadDeptEmail)
                                    CheckData = boolbox(msg="Is the next data correct?\n\n" + str(dataCheck), title=ProgramTitle)
                                    if CheckData == True:
                                        sendEmail(NameFile=NameFilePDF, FormDayDate=FileDay, ModifyFile = True, addr_from = "Company.ISM@Company.com", addr_to = HeadDeptEmail, addr_cc = "Company.ISM@Company.com", PdfFile = NameFilePDF1, SubjectMsg = SubjectMsg, SendSoxForm = True, UserName = userArrayToSend[x], StartDate = CorrectFormDate, Dept = None)
                                    else:
                                        continue
                        else:
                            continue       
                elif AccessOptions == "2":

                    userArray = []
                    headdeptArray = []
                    formdateArray = []
                    DeptArray = []
                    userArrayToSend = []
                    headdeptArrayToSend = []
                    formdateArrayToSend = []
                    DeptArrayToSend = []

                    NoData = 1
                    ModifyExcel = 0


                    for x in range(2,(Data.max_row)+1):
                        if Data.cell(row = x, column = 1).value == None or len(Data.cell(row = x, column = 1).value) < 1:
                            continue
                        else:
                            SendAnswer = Data.cell(row = x, column = 48).value
                            FirstName = Data.cell(row = x, column = 1).value
                            if charCounter(string = FirstName, char = " ") == 1:
                                FirstName = FirstName.replace(" ","_")
                                FirstName = (FirstName.split("_")[0]).capitalize() + "_" + (FirstName.split("_")[1]).capitalize()
                            LastName = Data.cell(row = x, column = 2).value
                            if charCounter(string = LastName, char = " ") == 1:
                                LastName = LastName.replace(" ","_")
                                LastName = (LastName.split("_")[0]).capitalize() + "_" + (LastName.split("_")[1]).capitalize()
                            Dept = Data.cell(row = x, column = 7).value
                            HeadDept1 = Data.cell(row = x, column = 8).value
                            FormDate = Data.cell(row = x, column = 10).value

                            UserNameToEnter = FirstName.replace("_"," ") + " " + LastName.replace("_"," ")
                            userArray.append(UserNameToEnter)
                            headdeptArray.append(HeadDept1)
                            formdateArray.append(FormDate)
                            DeptArray.append(Dept)
                    else:
                        userArrayToSend = multchoicebox(msg="Select Users to send the form", title=ProgramTitle, choices = userArray)
                        if len(userArrayToSend) != None:
                            headdeptArrayToSend[:] = headdeptArray
                            formdateArrayToSend[:] = formdateArray
                            DeptArrayToSend[:] = DeptArray
                            reverseuserArrayToSend = list(reversed(userArrayToSend))

                            for x in reverseuserArrayToSend:
                                UserIndex = userArray.index(x)
                                headdeptArray.remove(headdeptArrayToSend[UserIndex])
                                formdateArray.remove(formdateArrayToSend[UserIndex])
                                DeptArray.remove(DeptArrayToSend[UserIndex])
                            for x in range(0,len(headdeptArray)):
                                headdeptArrayToSend.remove(headdeptArray[x])
                                formdateArrayToSend.remove(formdateArray[x])
                                DeptArrayToSend.remove(DeptArray[x])
                            for x in range(0,len(userArrayToSend)):
                                dataCheck = []
                                CorrectFormDate = formdateArrayToSend[x].replace("/","-")
                                DateMForm = CorrectFormDate.split("-")[1]
                                DateYForm = CorrectFormDate.split("-")[2]

                                SubjectMsg = "Sox User Access Form " + months[int(DateMForm)] + " " + DateYForm + " - " + userArrayToSend[x].capitalize()
                                NameFilePDF = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] +  "\\Company_" + (userArrayToSend[x].capitalize()).replace(" ","_") + "-" + CorrectFormDate + ".pdf"
                                NameFileDocx = "Company_" + (userArrayToSend[x].capitalize()).replace(" ","_") + "-" + CorrectFormDate + ".docx"
                                NameFilePDF1 = NameFileDocx.replace(".docx",".pdf")

                                if charCounter(string = headdeptArrayToSend[x], char = " ") == 1:
                                    HeadDeptEmail = headdeptArrayToSend[x].replace(" ",".") + "@Company.com"
                                elif charCounter(string = headdeptArrayToSend[x], char = " ") == 2:
                                    HeadDeptEmail = headdeptArrayToSend[x].split(" ")[0] + headdeptArrayToSend[x].split(" ")[1] + "." + headdeptArrayToSend[x].split(" ")[2] + "@Company.com"

                                dataCheck.append(userArrayToSend[x])
                                dataCheck.append(formdateArrayToSend[x])
                                dataCheck.append(DeptArrayToSend[x])
                                dataCheck.append(HeadDeptEmail)
                                CheckData = boolbox(msg="Is the next data correct?\n\n" + str(dataCheck), title=ProgramTitle)
                                if CheckData == True:
                                    sendEmail(NameFile=NameFilePDF, FormDayDate=FileDay, ModifyFile = True, addr_from = "Company.ISM@Company.com", addr_to = HeadDeptEmail, addr_cc = "Company.ISM@Company.com", PdfFile = NameFilePDF1, SubjectMsg = SubjectMsg, SendSoxForm = True , Username = userArrayToSend[x], StartDate = CorrectFormDate, Dept = DeptArrayToSend[x])
                                else:
                                    continue
                        else:
                            continue
                elif AccessOptions == "3":

                    for x in range(2,Sox.max_row+1):
                        FirstName = Sox.cell(row = x, column = 1).value
                        LastName = Sox.cell(row = x, column = 2).value
                        EmployeeNum = Sox.cell(row = x, column = 6).value
                        if EmployeeNum == "XXXX" or EmployeeNum == "YYYY":
                            EmployeeNum = "-"
                        StartLeave = Sox.cell(row = x, column = 9).value
                        SLDate = Sox.cell(row = x, column = 10).value


                        if StartLeave == "Leaving":

                            DateYForm = SLDate.split("/")[2]
                            DateMForm = SLDate.split("/")[1]
                            CorrectFormDate = SLDate.replace("/","-")
                            os.makedirs(FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)], exist_ok=True)
                            os.makedirs(FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated", exist_ok=True)

                            if charCounter(string = FirstName, char = " ") == 1:
                                if charCounter(string = LastName, char = " ") == 1:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                elif charCounter(string = LastName, char = " ") == 2:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "-" + CorrectFormDate + ".docx"
                                elif charCounter(string = LastName, char = " ") == 3:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "_" + LastName.split(" ")[3] + "-" + CorrectFormDate + ".docx"
                                else:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + LastName + "-" + CorrectFormDate + ".docx"
                            elif charCounter(string = FirstName, char = " ") == 2:
                                if charCounter(string = LastName, char = " ") == 1:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                elif charCounter(string = LastName, char = " ") == 2:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "-" + CorrectFormDate + ".docx"
                                elif charCounter(string = LastName, char = " ") == 3:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "_" + LastName.split(" ")[3] + "-" + CorrectFormDate + ".docx"
                                else:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + LastName + "-" + CorrectFormDate + ".docx"
                            elif charCounter(string = FirstName, char = " ") == 3:
                                if charCounter(string = LastName, char = " ") == 1:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                elif charCounter(string = LastName, char = " ") == 2:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "-" + CorrectFormDate + ".docx"
                                elif charCounter(string = LastName, char = " ") == 3:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "_" + LastName.split(" ")[2] + "_" + LastName.split(" ")[3] + "-" + CorrectFormDate + ".docx"
                                else:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName.split(" ")[0] + "_" + FirstName.split(" ")[1] + "_" + FirstName.split(" ")[2] + "_" + FirstName.split(" ")[3] + "_" + LastName + "-" + CorrectFormDate + ".docx"

                            else:
                                if charCounter(string = LastName, char = " ") == 1:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName + "_" + LastName.split(" ")[0] + "_" + LastName.split(" ")[1] + "-" + CorrectFormDate + ".docx"
                                else:
                                    NameTerminatedFile = FormFiles + "\\" + DateYForm + "\\" + months[int(DateMForm)] + "\\Terminated" + "\\Company_Terminated_User_" + FirstName + "_" + LastName + "-" + CorrectFormDate + ".docx"

                            terminatingUsers(FirstName, LastName, EmployeeNum, SLDate, NameTerminatedFile)
                        else:
                            next
                elif AccessOptions == "4":
                    if optionModify == "1":
                        copySoxData(SoxFile =TMPSoxExcel, AccessFile = TMPAccessExcel)
                    if optionModify == "2":
                        copySoxData(SoxFile =TMPSoxExcel2, AccessFile = TMPAccessExcel2)
                    else:
                        copySoxData(SoxFile =TMPSoxExcel, AccessFile = TMPAccessExcel)

                elif AccessOptions == "5":
                    break                
        elif MainMenuChoice == "4":
            Exit = "Yes"
    else:
        msgbox(msg="File Created By Hector Rodriguez - " + DateVersion.split("/")[2],title=ProgramTitle,ok_button="Exit")
except:
	logging.exception(" Exception errors at " + time.strftime("%H:%M:%S") + "\n")

	with open("logs\\CreateSoxForm-" + time.strftime("%H%M-%A_%d_%B_%Y") + ".log") as LogFile:
		codebox(msg="There is an error on the program execution. Please send this to your IT Department", title=ProgramTitle,text=list(LogFile))